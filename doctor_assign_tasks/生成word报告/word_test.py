# coding=utf-8
import docx,copy
import pandas as pd
from docx import Document
from docx.oxml import CT_P,CT_Tbl
from docx.table import _Cell, Table, _Row
from docx.text.paragraph import Paragraph

def get_word(filename):             # 纯文字情况。
    document = docx.Document(filename)
    content = []
    for paragraph in document.paragraphs:
        print(paragraph.text)
        content.append(paragraph.text)
    return '\n\n'.join(content)

def get_all_table_text(path):
    document = Document(path)
    all_tables = document.tables
    text_list = []
    # for table in all_tables:
    for row in all_tables[0].rows:
        text = []
        for cell in row.cells:
            text.append(cell.text)
        text_list.append(text)
    print(text_list)

def machining_table(num_th):
    docx = Document('17.tmp.2023v1.docx')
    table = docx.tables[num_th]
    print(len(table.rows),len(table.columns))
    # name = table.cell(1,1).paragraphs[0].add_run('普阿梅')
    table.cell(1, 1).paragraphs[0].text = '普阿梅'
    table.cell(1, 3).paragraphs[0].text = '女'
    table.cell(2,3).paragraphs[0].text = '苏州大学附一医院'
    docx.save('new_17.tmp.2023v1.docx')


from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm
from docx.shared import RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
def inser_word_table():
    def move_table_after(table,paragraph):        # 从指定paragraph 后insert table。
        tbl, p = table._tbl, paragraph._p
        p.addnext(tbl)
    def Set_Background_Color(cell, rgbColor):
        shading_elm = parse_xml(
            r'<w:shd {} w:fill="{color_value}"/>'.format(nsdecls('w'), color_value=rgbColor))  # 固定写法，照抄即可
        cell._tc.get_or_add_tcPr().append(shading_elm)

    doc = Document('17.tmp.2023v1.docx')
    for paragraph in doc.paragraphs:
        paragraph_text = paragraph.text
        if paragraph_text.endswith('pathogenictable'):                  # 定位目标段落。
            target = paragraph
            print(target.text)
    table = doc.add_table(6, 3)
    table.style = 'Table Grid'
    run = table.cell(0,0).paragraphs[0].add_run('基因')
    table.cell(0,1).text, table.cell(0,2).text= '检测内容','检测结果（仅显示致病/可能致病变异）'
    table.cell(1, 0).text,table.cell(1, 1).text, table.cell(1, 2).text = 'ALK', '突变、重排','野生型/未检出'
    table.cell(2, 0).text, table.cell(2, 1).text, table.cell(2, 2).text = 'AKT1', '突变', '野生型/未检出'
    table.cell(3, 0).text, table.cell(3, 1).text, table.cell(3, 2).text = 'AKT1', '突变', '野生型/未检出'
    table.cell(4, 0).text, table.cell(4, 1).text, table.cell(4, 2).text = 'AKT1', '突变', '野生型/未检出'
    table.cell(5, 0).text, table.cell(5, 1).text, table.cell(5, 2).text = 'AKT1', '突变', '野生型/未检出'
    table.cell(0,0).merge(table.cell(1,0))                              # 合并单元格，也可以用对角线cell。
    table.cell(0,0).text = '基因'
    for j in range(len(table.columns)):                                 # 表格第一行 粗体。
        table.cell(0,j).paragraphs[0].runs[0].font.bold = True
    for i in range(len(table.rows)):                                    # 行距 1厘米。列宽调整。
        table.rows[i].height = Cm(1)
    col_width = [1,1,20]                                                # 按比例列宽调整。
    for col in range(3):
        table.cell(0,col).width = Cm(col_width[col])
    for i in range(0,1):                                                # 第一行背景色。
        for j in range(len(table.columns)):
            Set_Background_Color(table.cell(i,j), "FF0000")
    for row in table.rows:                                              # 调整字体。run就是text的格式信息。
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(13)
                    run.font.color.rgb = RGBColor(0, 0, 255)
                    run.font.name = 'Times New Roman'
                    from docx.oxml.ns import qn                         # 中文字体。
                    run._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
    for row in range(len(table.rows)):                                                # 居中
        for col in range(len(table.columns)):
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            table.cell(row,col).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER    # 水平居中
            from docx.enum.table import WD_ALIGN_VERTICAL
            table.cell(row, col).vertical_alignment = WD_ALIGN_VERTICAL.CENTER         # 垂直居中

    table_before_paragraph = target.insert_paragraph_before('Insert_Table')            # 插入段落。
    table_before_paragraph.runs[0].font.size = Pt(20)




    move_table_after(table,target)
    table1 = copy.deepcopy(doc.tables[2])                               # 深copy，因为python的list是可变对象。
    move_table_after(table1, target)

    def delete_paragraph(paragraph):                                    # 删除段落。
        p = paragraph._element
        p.getparent().remove(p)
        paragraph._p = paragraph._element = None
    delete_paragraph(target)
    doc.save('new_17.tmp.2023v1.docx')

def add_word_header():
    doc = Document('new_17.tmp.2023v1.docx')
    section = doc.sections[1]                                           # 在第二节添加页眉
    header = section.header
    doc.sections[2].header.is_linked_to_previous = False                # 为true说明这一节和上一节状态相同。
    doc.sections[1].header.is_linked_to_previous = False
    doc.sections[0].header.is_linked_to_previous = False

    doc.sections[1].header_distance = Cm(1)                             # 在第二节,设置页眉距离 往上1厘米。
    header.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER      # 居中显示
    header.paragraphs[0].text = '实验页眉'
    for run in header.paragraphs[0].runs:
        run.font.color.rgb = RGBColor(0x42, 0x24, 0xe9)
    print(len(doc.sections))
    for i in doc.sections:
        print(i.header.paragraphs[0].text)
    doc.save('new_17.tmp.2023v1.docx')

def Section_by_Contents():                                              # 按照目录要求分节，方便页眉按照节来定。
    pass



if __name__=='__main__':
    # print(get_word('17.tmp.2023v1.docx'))
    # get_all_table_text('17.tmp.2023v1.docx')
    # insert_table_after_text('17.tmp.2023v1.docx',"pathogenictable")
    # machining_table(0)

    inser_word_table()
    add_word_header()
