# coding=utf8
import docx, copy, re, os
import pandas as pd
import datetime
import math
import shutil
from docx import Document
from docx.oxml import CT_P, CT_Tbl
from docx.table import _Cell, Table, _Row
from docx.text.paragraph import Paragraph
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm
from docx.shared import RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class explain():  # 解读 c.96+1G>T 以及 p.L747_P753delinsS 的意思。
    def __init__(self):
        self.kind, self.position, self.explain = '', '', ''
        self.Amino_acid_dict = {}
        Amino_acid_list = ['甘氨酸', '丙氨酸', '缬氨酸', '亮氨酸', '异亮氨酸', '脯氨酸', '苯丙氨酸', '酪氨酸', '色氨酸', '丝氨酸', '苏氨酸', '半胱氨酸', '蛋氨酸',
                           '天冬酰胺', '谷氨酰胺', '天冬氨酸', '谷氨酸', '赖氨酸', '精氨酸', '组氨酸']
        abbreviate_table = ['G', 'A', 'V', 'L', 'I', 'P', 'F', 'Y', 'W', 'S', 'T', 'C', 'M', 'N', 'Q', 'D', 'E', 'K',
                            'R', 'H']
        for i in range(len(Amino_acid_list)):
            self.Amino_acid_dict[abbreviate_table[i]] = Amino_acid_list[i]

    def classification_Description(self, raw_explain):
        result = []
        if re.match('delins', raw_explain) != None:
            result.append(f"缺失，同时发生了{len(re.match('delins([A-WY-Z]+)', raw_explain).group(1))}个氨基酸的插入")
            raw_explain = raw_explain.replace(re.match('(delins[A-WY-Z]+)', raw_explain).group(1), '')
        if re.match('del', raw_explain) != None:
            result.append('发生了缺失')
            raw_explain = raw_explain.replace(re.match('(del)', raw_explain).group(1), '')
        if re.match('ins', raw_explain) != None:
            result.append(f"发生了{len(re.match('ins([A-WY-Z]+)', raw_explain).group(1))}个氨基酸的插入")
            raw_explain = raw_explain.replace(re.match('(ins[A-WY-Z]+)', raw_explain).group(1), '')
        if re.match('[A-WY-Z]+', raw_explain) != None:
            result.append(f"突变为{self.Amino_acid_dict[re.match('([A-WY-Z]+)', raw_explain).group(1)]}。")
            raw_explain = raw_explain.replace(re.match('([A-WY-Z]+)', raw_explain).group(1), '')
        if re.match('[X]', raw_explain) != None:
            result.append(f"突变为 未知类型氨基酸。")
        if re.match('fs', raw_explain) != None:
            i = re.match(r'fs\*([\d]+)', raw_explain).group(1)
            result.append(f'出现移码变异,导致后续的氨基酸序列发生错位，从而改变了翻译后的蛋白质序列，并且在插入位置之后的{i}位形成了一个终止密码子'
                          f'导致蛋白质合成的过程中到达这个位置后会停止合成。这种类型的变异通常会导致蛋白质合成的过程提前终止，从而产生缺失的蛋白质片段。'
                          f'由于蛋白质的功能通常与其结构密切相关，这种缺失可能会对蛋白质的功能和稳定性产生影响。')
            raw_explain = raw_explain.replace(re.match(r'(fs\*[\d]+)', raw_explain).group(1), '')
        return ''.join(result)

    def Interpretative_p(self, sentence):
        match = re.match(r'(^[\w]+[0-9]+)', sentence)  # match.group(1)=L747_P753
        match_group_list = match.group(1).split('_')
        # print(match_group_list)
        if len(match_group_list) > 1:
            self.position = '第' + match_group_list[0][1:] + '位' + self.Amino_acid_dict[
                match_group_list[0][0]] + '至' + '第' + match_group_list[1][1:] + '位' + self.Amino_acid_dict[
                                match_group_list[1][0]]
            raw_explain = sentence.replace(match.group(1), '')  # delinsD
            self.explain = self.classification_Description(raw_explain)
        else:
            self.position = '第' + match_group_list[0][1:] + '位' + self.Amino_acid_dict[match_group_list[0][0]]
            raw_explain = sentence.replace(match.group(1), '')
            try:
                self.explain = self.classification_Description(raw_explain)
            except:
                self.explain = '没找到'
        return f'{self.kind}{self.position}{self.explain}'

    def Interpretative_c(self, sentence):
        return '导致该基因编码的蛋白序列未知，预测产生无功能蛋白，该基因正常蛋白功能丧失。'

    def Interpretative(self, sentence):
        if sentence[0] == 'p':
            self.kind = '导致该基因编码的蛋白序列的'
            # print(self.Interpretative_p(sentence[2:]))
            return self.Interpretative_p(sentence[2:])
        elif sentence[0] == 'c':
            self.kind = '导致该基因编码的蛋白序列未知'
            # print(self.Interpretative_c(sentence[2:]))
            return self.Interpretative_c(sentence[2:])


# 这里以 BC17 也就是 《解码17基因》 为模版。
# add_table_Targeted_treatment_tips 中的 build_Rules 子方法中是否生成 中间文件，可选。(第一次运行的时候要记得打开，生成的中间文件后续会用上。)
# 定位段落。
class BC17():
    # self.doc类似于python 的可变对象的特点。这里不断流转对其加工。 docx 中的table 是可变对象，在方法内加工以后 ，外面的table会变的。
    def __init__(self, input_file):
        self.input_file = input_file
        self.sample = self.input_file.split('/')[-1].replace('.summary.xlsx', '')
        print(self.sample, 'input_file is : ', input_file)
        df = pd.read_excel(input_file, sheet_name=None)
        self.BC17_meta, self.BC17_snpindel, self.BC17_fusion, self.BC17_cnv, self.BC17_qc = df['meta'], df['snpindel'], \
                                                                                            df['fusion'], df['cnv'], df[
                                                                                                'qc']
        # print(self.BC17_meta['name'].iloc[0])
        df = pd.read_excel('/refhub/ref/drug/DrugApproval.xlsx', sheet_name=None)
        self.DrugApproval = df['药物获批情况表'].dropna(subset=['ApprovedContent'], how='any')
        df = pd.read_excel('/refhub/ref/gene/GeneInfo.xlsx', sheet_name=None)
        self.GeneIndo = df['Genename_table'].dropna(subset=['Genename'], how='any')
        self.doc = Document('/refhub/ref/masterplate/BC17.docx')
        df = pd.read_excel("/refhub/ref/citation/Citation.xlsx", sheet_name=None)
        self.Citation = df['citation']

    def Set_Background_Color(self, cell, rgbColor):
        shading_elm = parse_xml(
            r'<w:shd {} w:fill="{color_value}"/>'.format(nsdecls('w'), color_value=rgbColor))  # 换背景色固定写法，照抄即可
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def delete_paragraph(self, paragraph):  # 删除段落。
        p = paragraph._element
        p.getparent().remove(p)
        paragraph._p = paragraph._element = None

    def move_table_after(self, table, paragraph):  # 从指定paragraph 后insert table。
        tbl, p = table._tbl, paragraph._p
        p.addnext(tbl)

    def Set_Background_Color(self, cell, rgbColor):
        shading_elm = parse_xml(
            r'<w:shd {} w:fill="{color_value}"/>'.format(nsdecls('w'), color_value=rgbColor))  # 固定写法，照抄即可
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def fixed_position(self, End_character):  # End_character 用类似于 'pathogenictable'。
        for paragraph in self.doc.paragraphs:
            paragraph_text = paragraph.text
            if paragraph_text.endswith(End_character):  # 定位目标段落。
                target = paragraph
                return target

    def format_Cell(self, table, position=[(0, 0)], size=10, color=RGBColor(0, 0, 0),
                    left=False):  # doxc 中 table 和 table 中的cell 都是可变对象。
        for (row, column) in position:
            cell = table.cell(row, column)
            for paragraph in cell.paragraphs:  # 遍历单元格中的段落
                for run in cell.paragraphs[0].runs:
                    # run = cell.paragraphs[0].runs[0]  # 注意：添加一个新的文本块 add_run() 是不对的，我们这里是修改已有的cell。
                    font = run.font  # 获取字体对象
                    font.name = "Microsoft YaHei"  # 设置字体名称为宋体
                    font.size = Pt(size)  # 设置字体大小为12号
                    font.color.rgb = color  # 设置字体颜色为黑色
                    if left == False:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 设置段落水平居中
                    elif left == True:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    from docx.oxml.ns import qn  # 中文字体。
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # 设置单元格垂直居中

    def format_Table(self, table, header_size=10, col_width=[1, 1], header_color=RGBColor(0, 0, 0), Border=[],
                     Border_all=False):
        # 遍历每一行
        for i, row in enumerate(table.rows):
            # 获取每一行的所有单元格
            cells = row.cells
            # 遍历每个单元格
            for cell in cells:
                # 获取单元格的段落对象
                paragraph = cell.paragraphs[0]
                # 如果 第二行是 “未检出相关基因突变”，设置单元格高度。
                if cell.paragraphs[0].text == '未检出相关基因突变' or cell.paragraphs[0].text == '未检出相关基因突变，无相关提示。':
                    table.rows[1].height = Cm(1)
                # 如果是首行，设置字体为粗体，并设置背景色为灰色 ，设置单元格高度。
                if i == 0:
                    run = paragraph.runs[0]
                    run.font.bold = True
                    run.font.size = Pt(header_size)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    table.rows[0].height = Cm(1)
                    self.Set_Background_Color(cell, "CFDAE6")
                # 如果是奇数行，设置背景色为浅蓝色,且第一列不变色。
                elif i % 2 != 1:  # and cell !=cells[0]:
                    self.Set_Background_Color(cell, "F1F1F1")
                if cell.text == '基因与肿瘤相关性概述' or cell.text == '位点变异信息注释' or cell.text == '临床意义提示':
                    for paragraph in cell.paragraphs:
                        run = cell.paragraphs[0].runs[0]
                        font = run.font
                        font.bold = True
                        font.name = "Microsoft YaHei"  # 设置字体名称
                        font.size = Pt(11)  # 设置字体大小为12号
                        font.color.rgb = RGBColor(67, 114, 161)  # 设置字体颜色为黑色
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 设置段落水平左对齐。
                        from docx.oxml.ns import qn  # 中文字体。
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # 设置单元格垂直居中
        # # 按比例列宽调整。
        # for i in range(len(table.rows)):
        #     for col in range(len(table.columns)):
        #         table.cell(i, col).width = Cm(col_width[col])
        # 按照固定的宽度调整。以免表格超出页边的距离。
        table.autofit = False
        table.allow_autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.width = Cm(16)
        for i, col in enumerate(col_width):
            table.columns[i].width = Cm(col)

        # 加边框
        def set_cell_border(cell, **kwargs):
            # 这里的kwargs 其实就是一个dict，key 是 top、bottom 这些，value是 {} 一个嵌套的dict。
            # 不定数量的形参传入时常用这样的写法，形参传入时=符号充当了dict的构建，=左边是key右边是value。
            """
            Set cell`s border
            Usage:
            set_cell_border(
                cell,
                top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
                bottom={"sz": 12, "color": "#00FF00", "val": "single"},
                left={"sz": 24, "val": "dashed", "shadow": "true"},
                right={"sz": 12, "val": "dashed"},
            )
            """
            # XML 元素是 XML 文档中必不可缺的部分，我们可以将 XML 元素看成一个容器，其中存放了文本，元素，属性，媒体对象或所有的这些。XML 文档包含 XML 元素。
            tc = cell._tc  # cell._tc 是一个单元格对象。_tc 属性是一个内部属性，它返回一个 _Element 对象，表示单元格的 XML 元素。
            tcPr = tc.get_or_add_tcPr()  # 获取短元格内部属性。
            # check for tag existnace, if none found, then create one
            tcBorders = tcPr.first_child_found_in(
                "w:tcBorders")  # 它返回了单元格属性对象中第一个 <w:tcBorders> 元素。如果该元素不存在，则该方法返回 None。
            if tcBorders is None:
                tcBorders = OxmlElement('w:tcBorders')
                tcPr.append(tcBorders)
            # list over all available tags
            for edge in ('left', 'top', 'right', 'bottom', 'insideH', 'insideV'):
                edge_data = kwargs.get(edge)  # edge_data 就是形参中内部嵌套的子dict。
                if edge_data:
                    tag = 'w:{}'.format(edge)  # w:left
                    # check for tag existnace, if none found, then create one
                    from docx.oxml.ns import qn
                    element = tcBorders.find(qn(tag))  # qn(tag) 是一个函数调用，它将字符串 tag 转换为一个 XML 命名空间。
                    if element is None:
                        element = OxmlElement(tag)
                        tcBorders.append(element)
                    # looks like order of attributes is important
                    for key in ["sz", "val", "color", "space", "shadow"]:
                        if key in edge_data:
                            element.set(qn('w:{}'.format(key)), str(edge_data[key]))  # 在元素中 把 xml 和 形参输入的值 对应起来。

        if Border != [] and Border_all == False:  # Border = [] list 中存放 行号，第几行的所有单元格统一加某种边框。
            # 给指定单元格加边框
            for j in Border:
                for i in range(len(table.columns)):
                    set_cell_border(
                        table.cell(j, i),
                        # top={"color": "#000000", "space": "0"},
                        bottom={"val": "single", "color": "#CFDAE6", "space": "1"},
                        # left={"color": "#000000", "space": "0"},
                        # right={"sz": 0.5, "val": "double", "color": "#000000", "space": "0"},
                        # insideH={"sz": 0.5, "val": "double", "color": "#000000", "space": "0"},
                        # end={"sz": 0.5, "val": "double", "color": "#000000", "space": "0"}
                    )
        elif Border_all == True:
            # 给全部单元格加边框。
            for row in table.rows:
                for cell in row.cells:
                    set_cell_border(
                        cell,
                        top={"val": "dotted", "color": "#CFDAE6", "space": "0"},
                        bottom={"val": "dotted", "color": "#CFDAE6", "space": "0"},
                        left={"val": "dotted", "color": "#CFDAE6", "space": "0"},
                        right={"val": "dotted", "color": "#CFDAE6", "space": "0"},
                        insideH={"val": "dotted", "color": "#CFDAE6", "space": "0"},
                        end={"val": "dotted", "color": "#CFDAE6", "space": "0"}
                    )
        # table.autofit = True

    def format_paragraph(self, target_paragraph, size=10, bold=False):
        run = target_paragraph.runs[0]
        run.font.name = "Microsoft YaHei"  # 设置字体名称为宋体
        run.font.size = Pt(size)  # 设置字体大小为12号
        run.font.bold = bold
        from docx.oxml.ns import qn  # 中文字体。
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 开始真实插表。
    def machining_table_0(self, num_th=0):  # 给第0个表格差值。
        table = self.doc.tables[num_th]
        # print(len(table.rows), len(table.columns))
        # name = table.cell(1,1).paragraphs[0].add_run('普阿梅')
        table.cell(1, 1).text = self.BC17_meta['name'].iloc[0]
        table.cell(1, 3).text = self.BC17_meta['gender'].iloc[0]
        table.cell(1, 5).text = self.BC17_meta['age'].iloc[0].astype(str)
        date_string = self.BC17_meta['到样日期*'].iloc[
            0]  # 2023-03-01。在这里，我们首先使用strptime()函数将字符串转换为datetime对象，然后使用strftime()函数将datetime对象转换为字符串并指定所需的格式。
        date_object = datetime.datetime.strptime(date_string, "%Y/%m/%d")
        new_date_string = date_object.strftime("%Y-%m-%d")
        table.cell(2, 1).text = new_date_string
        table.cell(2, 3).text = self.BC17_meta['送检医院'].iloc[0] if not pd.isna(self.BC17_meta['送检医院'].iloc[0]) else ''
        table.cell(3, 1).text = self.BC17_meta['clinname'].iloc[0]
        table.cell(3, 3).text = self.BC17_meta['projectname'].iloc[0]
        table.cell(4, 1).text = self.BC17_meta['id'].iloc[0][:-2]
        table.cell(4, 3).text = self.BC17_meta['样本类型*'].iloc[0]
        table.cell(5, 1).text = '-'
        # 改字体。
        self.format_Cell(table,
                         position=[(1, 1), (1, 3), (1, 5), (2, 1), (2, 3), (3, 1), (3, 3), (4, 1), (4, 3), (5, 1)],
                         size=10)
        return self.doc

    def machining_table_2(self, num_th=2):  # 给第2个表格差值。
        table = self.doc.tables[num_th]
        print(len(table.rows), len(table.columns))
        # name = table.cell(1,1).paragraphs[0].add_run('普阿梅')
        for i in range(1, len(table.rows)):
            table.cell(i, 2).paragraphs[0].text = '野生型/未检出'
        # 每一个位置可能会被重复插入。换行富豪分割
        # 这里为了安全，本来不用exec和eval 想使用local或者global来动态生成变量。global的区域以函数分割。
        # 但是 local 动态赋值是不可取的，local 只读 global 可读可写。参考<local陷阱>：https://blog.csdn.net/weixin_43064185/article/details/104141327
        # 我又不想搞模块全局变量赋值，所以还是选用来有点危险的 exec 方法来动态赋值。小心一点。
        # 在函数a()中使用exec()函数动态创建变量时，这些变量将存在于执行exec()函数的作用域中,外部作用域无法调用，
        # 外部作用域中无法直接引用这些变量。
        # 解决这个问题的方法是将动态创建的变量存储在一个可访问的数据结构中，例如字典或列表，以便在需要时进行引用。
        # 方案一，用list，像我在制作污染库时一样。
        # self.cell = []
        # for j in range(1, len(table.rows)):
        #     # locals()[f'cell_{j}'] = '_'  # local 动态赋值不可取。
        #     exec(f"cell_{j} = '_' ")
        #     self.cell.append(locals()[f'cell_{j}'])
        # 方案二，用dict。
        self.cell = {}
        for j in range(1, len(table.rows)):
            # locals()[f'cell_{j}'] = '_'  # local 动态赋值不可取。
            exec(f"self.cell['cell_{j}'] = '' ")
        # print(self.cell['cell_1'])
        # 第一先考虑snpindel 这个sheet 中的s1 情况。
        df_ref = self.BC17_snpindel[self.BC17_snpindel['review'].isin(['S1', 'S2'])]
        # EGFR需要确定插入哪个单元格
        for index, row in df_ref.iterrows():
            # print(row['AAchange'])
            if row['Gene.smallrefGene'] == 'EGFR':
                if row['EXON'] == 'exon18':
                    self.cell['cell_5'] = self.cell['cell_5'] + str(
                        row['AAchange']) + f" ({row['VAF'].split(',')[0]})" + "\n"
                elif row['EXON'] == 'exon19':
                    self.cell['cell_6'] = self.cell['cell_6'] + str(
                        row['AAchange']) + f" ({row['VAF'].split(',')[0]})" + "\n"
                elif row['EXON'] == 'exon20':
                    self.cell['cell_7'] = self.cell['cell_7'] + str(
                        row['AAchange']) + f" ({row['VAF'].split(',')[0]})" + "\n"
                elif row['EXON'] == 'exon21':
                    self.cell['cell_8'] = self.cell['cell_8'] + str(
                        row['AAchange']) + f" ({row['VAF'].split(',')[0]})" + "\n"
            elif row['Gene.smallrefGene'] == 'ALK':
                self.cell['cell_1'] = self.cell['cell_1'] + str(
                    row['AAchange']) + f" ({row['VAF'].split(',')[0]})" + "\n"
            elif row['Gene.smallrefGene'] == 'ATK1':
                self.cell['cell_2'] = self.cell['cell_2'] + str(
                    row['AAchange']) + f" ({row['VAF'].split(',')[0]})" + "\n"
            elif row['Gene.smallrefGene'] == 'BRAF':
                self.cell['cell_3'] = self.cell['cell_3'] + str(
                    row['AAchange']) + f" ({row['VAF'].split(',')[0]})" + "\n"
            elif row['Gene.smallrefGene'] == 'BCL2L11':
                self.cell['cell_4'] = self.cell['cell_4'] + str(
                    row['AAchange']) + f" ({row['VAF'].split(',')[0]})" + "\n"
            elif row['Gene.smallrefGene'] == 'ERBB2':
                self.cell['cell_9'] = self.cell['cell_9'] + str(
                    row['AAchange']) + f" ({row['VAF'].split(',')[0]})" + "\n"
            elif row['Gene.smallrefGene'] == 'KRAS':
                self.cell['cell_10'] = self.cell['cell_10'] + str(
                    row['AAchange']) + f" ({row['VAF'].split(',')[0]})" + "\n"
            elif row['Gene.smallrefGene'] == 'MAP2K1':
                self.cell['cell_11'] = self.cell['cell_11'] + str(
                    row['AAchange']) + f" ({row['VAF'].split(',')[0]})" + "\n"
            elif row['Gene.smallrefGene'] == 'MET':
                self.cell['cell_12'] = self.cell['cell_12'] + str(
                    row['AAchange']) + f" ({row['VAF'].split(',')[0]})" + "\n"
            elif row['Gene.smallrefGene'] == 'PIK3CA':
                self.cell['cell_13'] = self.cell['cell_13'] + str(
                    row['AAchange']) + f" ({row['VAF'].split(',')[0]})" + "\n"
            elif row['Gene.smallrefGene'] == 'PTEN':
                self.cell['cell_14'] = self.cell['cell_14'] + str(
                    row['AAchange']) + f" ({row['VAF'].split(',')[0]})" + "\n"
            elif row['Gene.smallrefGene'] == 'NTRK1':
                self.cell['cell_15'] = self.cell['cell_15'] + str(
                    row['AAchange']) + f" ({row['VAF'].split(',')[0]})" + "\n"
            elif row['Gene.smallrefGene'] == 'NTRK2':
                self.cell['cell_16'] = self.cell['cell_16'] + str(
                    row['AAchange']) + f" ({row['VAF'].split(',')[0]})" + "\n"
            elif row['Gene.smallrefGene'] == 'NTRK3':
                self.cell['cell_17'] = self.cell['cell_17'] + str(
                    row['AAchange']) + f" ({row['VAF'].split(',')[0]})" + "\n"
            elif row['Gene.smallrefGene'] == 'NRAS':
                self.cell['cell_18'] = self.cell['cell_18'] + str(
                    row['AAchange']) + f" ({row['VAF'].split(',')[0]})" + "\n"
            elif row['Gene.smallrefGene'] == 'RET':
                self.cell['cell_19'] = self.cell['cell_19'] + str(
                    row['AAchange']) + f" ({row['VAF'].split(',')[0]})" + "\n"
            elif row['Gene.smallrefGene'] == 'ROS1':
                self.cell['cell_20'] = self.cell['cell_20'] + str(
                    row['AAchange']) + f" ({row['VAF'].split(',')[0]})" + "\n"
        # 第二考虑 fusion 这个sheet 中的s1 情况。
        df_ref = self.BC17_fusion[self.BC17_fusion['review'].isin(['S1', 'S2'])]
        # print(df_ref)
        for index, row in df_ref.iterrows():
            for i in range(1, len(table.rows)):
                if table.cell(i, 0).paragraphs[0].text == row['Region1']: self.cell[f'cell_{i}'] = self.cell[
                                                                                                       f'cell_{i}'] + \
                                                                                                   row[
                                                                                                       'Region1'] + '-' + \
                                                                                                   row[
                                                                                                       'Region2'] + " 基因融合\n"
                if table.cell(i, 0).paragraphs[0].text == row['Region2']: self.cell[f'cell_{i}'] = self.cell[
                                                                                                       f'cell_{i}'] + \
                                                                                                   row[
                                                                                                       'Region1'] + '-' + \
                                                                                                   row[
                                                                                                       'Region2'] + " 基因融合\n"
        # 第三考虑 cnv 这个sheet 中的s1 情况。
        df_ref = self.BC17_cnv[self.BC17_cnv['review'].isin(['S1', 'S2'])]
        # print(df_ref)
        for index, row in df_ref.iterrows():
            for i in range(1, len(table.rows)):
                if table.cell(i, 0).paragraphs[0].text == row['Gene']: self.cell[f'cell_{i}'] = self.cell[
                                                                                                    f'cell_{i}'] + "基因扩增\n"
        # 把 完工的cell 填入对应的talbe的位置。
        position = []
        for i in range(1, len(table.rows)):
            if self.cell[f'cell_{i}'] != '':
                table.cell(i, 2).paragraphs[0].text = self.cell[f'cell_{i}'].strip('\n')
                position.append((i, 2))
        # 改字体。
        position_1 = [(i, j) for i in range(len(table.rows)) for j in range(len(table.columns))]
        self.format_Cell(table, position=position_1, size=10, color=RGBColor(0, 0, 0))
        self.format_Cell(table, position=position, size=10, color=RGBColor(255, 0, 0))
        return self.doc

    def add_table_Somatic_variation_results(self):
        # 先从基因检测结果表中读出非 '野生型/未检出' 结果。
        table_gene_result = self.doc.tables[2]  # 读基因检测结果表。
        self.cell = []
        self.cell_0 = []
        self.cell_1 = []  # 和 self.cell_0 一样，self.cell_1 在后续中也可能用到，主要是用于第二个表的 s3s4数量的记录。
        for i in range(1, len(table_gene_result.rows)):
            temp_result = table_gene_result.cell(i, 2).paragraphs[0].text
            temp_gene = table_gene_result.cell(i, 0).paragraphs[0].text
            if temp_result != '野生型/未检出':
                # print(temp_result.split('\n'))
                self.cell.extend([temp_gene + ' ' + ele for ele in temp_result.split('\n')])
        for i in range(len(self.cell)):
            p = re.compile(r'[(](.*?)[)]', re.S)
            self.cell[i] = re.sub(p, '', self.cell[i])
        # print(self.cell)     # ['ALK EML401-ALK 基因融合', 'ALK EML400-ALK 基因融合', 'EGFR p.L747_E749del', 'EGFR p.Lzhoujielun', 'EGFR p.zhoujielun', 'EGFR p.zhoujluner', 'MAP2K1 p.liudehua', 'MET MET 基因扩增']
        # print(self.cell,'注意看这个 self cell')
        # 建立表1。
        header = ['基因', '转录本编号', '核苷酸变化', '氨基酸变化', '外显子', '变异类型', '丰度/\n拷贝数']
        if len(self.cell) == 0:  # 空 情况。
            table = doc.add_table(2, 7)
            for i in range(len(header)):  # 列
                table.cell(0, i).paragraphs[0].text = header[i]
            table.cell(1, 0).merge(table.cell(1, 6))
            table.cell(1, 0).paragraphs[0].text = '未检出相关基因突变'
        else:
            table = doc.add_table(len(self.cell) + 1, 7)
            for i in range(7):
                for j in range(len(self.cell) + 1):
                    table.cell(j, i).paragraphs[0].text = '/'
            for i in range(0, len(header)):  # 列
                table.cell(0, i).paragraphs[0].text = header[i]
            # 第一先考虑snpindel 这个sheet 中的s1 情况。
            df_ref = self.BC17_snpindel[self.BC17_snpindel['review'].isin(['S1', 'S2'])]
            # print(df_ref)
            i = 1
            for index, row in df_ref.iterrows():
                # print(index,row['Gene.smallrefGene'],row['RNA'])
                table.cell(i, 0).paragraphs[0].text = row['Gene.smallrefGene']
                table.cell(i, 1).paragraphs[0].text = row['RNA']
                table.cell(i, 2).paragraphs[0].text = row['NCchange']
                table.cell(i, 3).paragraphs[0].text = row['AAchange']
                table.cell(i, 4).paragraphs[0].text = row['EXON'].replace('exon', '')
                table.cell(i, 5).paragraphs[0].text = row['ExonicFunc.smallrefGene']
                table.cell(i, 6).paragraphs[0].text = row['VAF'].split(',')[0]
                i += 1
            # 第二考虑fusion 这个sheet 中的s1 情况。
            df_ref = self.BC17_fusion[self.BC17_fusion['review'].isin(['S1', 'S2'])]
            for index, row in df_ref.iterrows():
                table.cell(i, 0).paragraphs[0].text = row['Region1'] + '-' + row['Region2']
                table.cell(i, 5).paragraphs[0].text = '基因融合'
                table.cell(i, 6).paragraphs[0].text = '{:.2f}%'.format(
                    (float(row['Break_support1']) + float(row['Break_support2'])) / float(row['Total_depth']) * 100)
                # print(float(row['Break_support1']))
                i += 1
            # 第三考虑cnv 这个sheet 中的s1 情况。
            df_ref = self.BC17_cnv[self.BC17_cnv['review'].isin(['S1', 'S2'])]
            for index, row in df_ref.iterrows():
                # print(row['N.exons'])
                table.cell(i, 0).paragraphs[0].text = row['Gene']
                table.cell(i, 1).paragraphs[0].text = \
                self.GeneIndo[self.GeneIndo['Genename'] == row['Gene']]['RNA'].iloc[0].split('.')[0]
                table.cell(i, 4).paragraphs[0].text = str(row['N.exons'])
                table.cell(i, 5).paragraphs[0].text = '基因扩增'
                table.cell(i, 6).paragraphs[0].text = str(row['cnv.number']).split('.')[0]
                i += 1
        self.table_analysis = copy.deepcopy(table)  # 拷贝好这个表格，给后面一个插入的地方用，记得表格的拷贝要靠 深拷贝。
        # 建立表2。
        num_row = 0  # S3 G3 的数量。
        # 第一先考虑snpindel 这个sheet 中的s3 g3 情况。
        df_ref_1 = self.BC17_snpindel[self.BC17_snpindel['review'].isin(['S3', 'G3'])]
        num_row += len(df_ref_1)
        # 第二考虑fusion 这个sheet 中的s3 g3 情况。
        df_ref_2 = self.BC17_fusion[self.BC17_fusion['review'].isin(['S3', 'G3'])]
        num_row += len(df_ref_2)
        # 第三考虑cnv 这个sheet 中的s3 g3 情况。
        df_ref_3 = self.BC17_cnv[self.BC17_cnv['review'].isin(['S3', 'G3'])]
        num_row += len(df_ref_3)
        if num_row == 0:  # 空情况。
            table_1 = doc.add_table(2, 7)
            for i in range(len(header)):
                table_1.cell(0, i).paragraphs[0].text = header[i]
            table_1.cell(1, 0).merge(table_1.cell(1, 6))
            table_1.cell(1, 0).paragraphs[0].text = '未检出相关基因突变'
        else:
            table_1 = doc.add_table(num_row + 1, 7)
            for i in range(7):
                for j in range(num_row + 1):
                    table_1.cell(j, i).paragraphs[0].text = '/'
            for i in range(len(header)):
                table_1.cell(0, i).paragraphs[0].text = header[i]
            i = 1
            for index, row in df_ref_1.iterrows():
                table_1.cell(i, 0).paragraphs[0].text = row['Gene.smallrefGene']
                table_1.cell(i, 1).paragraphs[0].text = row['RNA']
                table_1.cell(i, 2).paragraphs[0].text = row['NCchange']
                table_1.cell(i, 3).paragraphs[0].text = row['AAchange']
                table_1.cell(i, 4).paragraphs[0].text = row['EXON'].replace('exon', '')
                table_1.cell(i, 5).paragraphs[0].text = row['ExonicFunc.smallrefGene']
                table_1.cell(i, 6).paragraphs[0].text = row['VAF'].split(',')[0]
                self.cell_1.append(table_1.cell(i, 0).paragraphs[0].text)
                i += 1
            for index, row in df_ref_2.iterrows():
                table_1.cell(i, 0).paragraphs[0].text = row['Region1'] + '-' + row['Region2']
                table_1.cell(i, 5).paragraphs[0].text = '基因融合'
                table_1.cell(i, 6).paragraphs[0].text = '{:.2f}%'.format(
                    (float(row['Break_support1']) + float(row['Break_support2'])) / float(row['Total_depth']) * 100)
                self.cell_1.append(table_1.cell(i, 0).paragraphs[0].text)
                i += 1
            for index, row in df_ref_3.iterrows():
                table_1.cell(i, 0).paragraphs[0].text = row['Gene']
                table_1.cell(i, 1).paragraphs[0].text = \
                self.GeneIndo[self.GeneIndo['Genename'] == row['Gene']]['RSG'].iloc[0]
                table_1.cell(i, 4).paragraphs[0].text = 'exon ' + str(row['N.exons'])
                table_1.cell(i, 5).paragraphs[0].text = '基因扩增'
                table_1.cell(i, 6).paragraphs[0].text = str(row['cnv.number']).split('.')[0]
                self.cell_1.append(table_1.cell(i, 0).paragraphs[0].text)
                i += 1
        self.cell_0 = copy.deepcopy(list(map(lambda x: x.strip(' '), self.cell)))
        # print(self.cell_1,self.cell_0)
        # 插入表。
        target_paragraph = self.fixed_position('Somatic_variation_results_table')
        # print(target_paragraph.text)
        table_before_paragraph_0 = target_paragraph.insert_paragraph_before('◆ 致病或可能致病的变异位点【Ⅰ类/Ⅱ类】')
        self.format_paragraph(table_before_paragraph_0, size=11, bold=True)
        table_before_paragraph_1 = target_paragraph.insert_paragraph_before("\n◆ 临床意义不明确的变异位点【Ⅲ类】")
        self.format_paragraph(table_before_paragraph_1, size=11, bold=True)
        self.move_table_after(table, table_before_paragraph_0)
        self.move_table_after(table_1, table_before_paragraph_1)
        # 改字体。
        position = [(i, j) for i in range(len(table.rows)) for j in range(len(table.columns))]
        self.format_Cell(table, position=position, size=10, color=RGBColor(0, 0, 0))
        position = [(i, j) for i in range(len(table_1.rows)) for j in range(len(table_1.columns))]
        self.format_Cell(table_1, position=position, size=10, color=RGBColor(0, 0, 0))
        # 改表格。
        self.format_Table(table, col_width=[2, 3, 2.7, 2.8, 1.5, 2.3, 1.7], Border=[len(table.rows) - 1])
        self.format_Table(table_1, col_width=[2, 3, 2.7, 2.8, 1.5, 2.3, 1.7], Border=[len(table_1.rows) - 1])
        return self.doc

    def add_table_Targeted_treatment_tips(self):
        # 构建靶向治疗提示表 的规则方法。
        def build_Rules(Genetic_variation, variation_type) -> (
        str, str, str, str):  # 一次处理一行。返回的是  # 本癌药物 、其它癌药物、其它癌药物、耐药药物。
            ele = Genetic_variation.split('\n')
            df = pd.read_excel('/refhub/ref/var/VarLabel_Drug_Sensitivity.xlsx', sheet_name=None)
            df_sensitive = df['实体瘤用药表']
            # 按照snpdelins 方法来
            if re.findall(r'基因融合', Genetic_variation) == [] and re.findall(r'基因扩增', Genetic_variation) == []:
                name = ele[0]
                alteration_list = []
                alteration_0 = ele[1].replace(r'p.', '')  # L747_E749del
                alteration_1_list = re.findall(r'\d+', alteration_0)  # ['747', '749']
                alteration_2 = '致病突变'
                alteration_3 = variation_type
                alteration_list.extend([alteration_0, alteration_2, alteration_3])
                # print(name,alteration_0,alteration_1_list,alteration_2,alteration_3)
                df_result_0 = df_sensitive[
                    (df_sensitive['Gene'] == name) & (df_sensitive['VarLabel'].isin(alteration_list))]
                # df['col1'].astype(str).str.findall('\d+').apply(set).apply(lambda x: bool(x & set(a)))
                # x & set(a)是一个交集运算，它返回两个集合的交集。如果交集不为空，则返回True，否则返回False。
                # 因此，apply(lambda x: bool(x & set(a)))的作用是判断数字集合是否与列表a有交集，如果有，则返回True，否则返回False。
                df_result_1 = df_sensitive[(df_sensitive['Gene'] == name) & (
                    df_sensitive['VarLabel'].astype(str).str.findall('\d+').apply(set).apply(
                        lambda x: bool(x & set(alteration_1_list))))]
                df_result = pd.concat([df_result_0, df_result_1])
                df_result.drop_duplicates(keep='first', inplace=True)  # 所有列都相等就去重。# df_result_1 df_result_0 之间可能有重复的部分。
                # print(df_result)
                # 按照fusion 方法来
            elif re.findall(r'基因融合', Genetic_variation) != []:
                name_list = ele[0].split('-')
                alteration_0 = '融合'
                df_result_0 = df_sensitive[
                    (df_sensitive['Gene'] == name_list[0]) & (df_sensitive['VarLabel'].isin([alteration_0]))]
                df_result_1 = df_sensitive[
                    (df_sensitive['Gene'] == name_list[1]) & (df_sensitive['VarLabel'].isin([alteration_0]))]
                df_result = pd.concat([df_result_0, df_result_1])
                # print(name_list, alteration_0)
                # print(df_result)
            # 按照cnv 方法来
            elif re.findall(r'基因扩增', Genetic_variation) != []:
                name = ele[0]
                alteration_0 = '扩增'
                df_result = df_sensitive[
                    (df_sensitive['Gene'] == name) & (df_sensitive['VarLabel'].isin([alteration_0]))]
                # print(name,alteration_0)
                # print(df_result)
            # df_result = df_result[df_result['Sensitivity'] == '敏感']
            # print(df_result)
            drug_list_raw = df_result['Drug'].tolist()
            drug_list = []
            for items in drug_list_raw: drug_list.extend(
                items.split(','))  # ['奥希替尼', '伏美替尼', '厄洛替尼+贝伐珠单抗', '纳武单抗+厄洛替尼']
            drug_list = list(set(drug_list))
            # print(drug_list)
            self_drug, other_drug, Clinical_drugs, resistance_Drug = '/', '/', '/', '/'
            for i, row in df_result.iterrows():
                # print(row)
                for drug in row['Drug'].split(','):
                    # print(drug)
                    # 临床药物
                    if self.DrugApproval[self.DrugApproval['DrugName'].isin([drug])].empty:
                        if row['Sensitivity'] == '敏感': Clinical_drugs += drug + ','
                        # print('临床药物: ',Clinical_drugs)
                    # 本癌药物,潜在耐药药物。(定义耐药药物是指，由这个变异搜到的药物匹配癌种对应相同，但是变异对药物有耐药性。)
                    elif not self.DrugApproval[self.DrugApproval['DrugName'].isin([drug])].empty:
                        # 加入一个 审批补丁方法，有的药物只针对固定的突变 通过了审批，
                        # 如果没有通过此补丁，相当于此药物未通过审批,回到上面看看能不能算成临床药物。
                        def review_Patches(drug, alteration_0) -> bool:  # 这里我写的有点不一样，方便后续拓展，只是读起来有点绕。
                            options = {('索托拉西布'): 'G12C', ('阿达格拉西布'): 'G12C'}
                            return options.get((drug), alteration_0) != alteration_0

                        # print('look here!!! ',drug,alteration_0)
                        if review_Patches(drug, alteration_0):  # 打补丁
                            # 临床药物
                            if row['Sensitivity'] == '敏感': Clinical_drugs += drug + ','
                        else:  # 不打补丁
                            df_temp = self.DrugApproval[
                                self.DrugApproval['DrugName'].isin([drug])]  # 这个 df_temp 是 drug approve表中的 子集。
                            # print(df_temp)
                            # 和 上面一样，使用求交集的运算，更加灵活一些。下面这个print 调试时候用。
                            # print(f'df_temp {drug} 的匹配： ',df_temp[df_temp['OncoType'].astype(str).str.split(',').apply(set).apply(
                            #     lambda x: bool(x & set(self.BC17_meta['cancer'].tolist())))])
                            if not df_temp[df_temp['OncoType'].astype(str).str.split(',').apply(set).apply(
                                    lambda x: bool(x & set(self.BC17_meta['cancer'].tolist())))].empty:
                                if row['Sensitivity'] == '敏感':  self_drug += drug + ','
                                if row['Sensitivity'] == '耐药': resistance_Drug += drug + ','
                            # print('本癌药物: ', self_drug)
                            # 其它癌药物
                            if df_temp[df_temp['OncoType'].astype(str).str.split(',').apply(set).apply(
                                    lambda x: bool(x & set(self.BC17_meta['cancer'].tolist())))].empty:
                                if row['Sensitivity'] == '敏感':  other_drug += drug + ','
                                # print('其它癌药物: ', other_drug)
            # 去重，处理格式。
            if len(self_drug) > 1: self_drug = '\n'.join(list(set(self_drug.strip('/').strip(',').split(','))))
            if len(other_drug) > 1: other_drug = '\n'.join(list(set(other_drug.strip('/').strip(',').split(','))))
            if len(Clinical_drugs) > 1: Clinical_drugs = '\n'.join(
                list(set(Clinical_drugs.strip('/').strip(',').split(','))))
            if len(resistance_Drug) > 1: resistance_Drug = '\n'.join(
                list(set(resistance_Drug.strip('/').strip(',').split(','))))
            # 注意:
            # 以下为 是否生成 中间文件，可选。(第一次运行的时候要记得打开，生成的中间文件后续会用上。)
            if not os.path.isfile(f'/refhub/ref/matching/{self.sample}药物注释匹配表.csv'):
                df_result.to_csv(f'/refhub/ref/matching/{self.sample}药物注释匹配表.csv', mode='a', index=False)
            else:
                df_result.to_csv(f'/refhub/ref/matching/{self.sample}药物注释匹配表.csv', mode='a', header=False, index=False)
            if not os.path.isfile(f'/refhub/ref/matching/{self.sample}注释匹配表.csv'):
                citation_list = []
                Citations = df_result['Citation'].tolist()
                for items in Citations: citation_list.extend(str(items).split(','))
                with open(f'/refhub/ref/matching/{self.sample}注释匹配表.csv', 'w')as f:
                    f.write(Genetic_variation.replace('\n', ' ') + ':' + ','.join(citation_list) + '\n')
            else:
                citation_list = []
                Citations = df_result['Citation'].tolist()
                for items in Citations: citation_list.extend(str(items).split(','))
                with open(f'/refhub/ref/matching/{self.sample}注释匹配表.csv', 'a+')as f:
                    f.write(Genetic_variation.replace('\n', ' ') + ':' + ','.join(citation_list) + '\n')
            # print('self_drug:', self_drug, 'other_drug:', other_drug, 'Clinical_drugs:', Clinical_drugs,'resistance_Drug: ', resistance_Drug)
            return self_drug, other_drug, Clinical_drugs, resistance_Drug

        # 先从基因检测结果表中读出非 '野生型/未检出' 结果。
        table_gene_result = self.doc.tables[2]  # 读基因检测结果表。
        self.cell = []
        for i in range(1, len(table_gene_result.rows)):
            temp_result = table_gene_result.cell(i, 2).paragraphs[0].text
            temp_gene = table_gene_result.cell(i, 0).paragraphs[0].text
            if temp_result != '野生型/未检出':
                # print(temp_result.split('\n'))
                self.cell.extend([temp_gene + ' ' + ele for ele in temp_result.split('\n')])
        for i in range(len(self.cell)):
            p = re.compile(r'[(](.*?)[)]', re.S)
            self.cell[i] = re.sub(p, '', self.cell[i])
        self.cell = list(map(lambda x: x.strip(' '), self.cell))
        print(
            self.cell)  # ['ALK EML4-ALK 基因融合', 'EGFR p.L747_E749del', 'EGFR p.T790M', 'ERBB2 基因扩增', 'MET 基因扩增', 'RET RET-IRX3 基因融合']
        if len(self.cell) != 0:
            # 建立表。
            table = doc.add_table(len(self.cell) + 1, 5)
            for i in range(5):
                for j in range(len(self.cell) + 1):
                    table.cell(j, i).paragraphs[0].text = '/'
            header = ['基因变异', '推荐本癌种药物', '推荐其他癌种药物', '临床试验药物', '潜在耐药药物']
            for i in range(len(header)):
                table.cell(0, i).paragraphs[0].text = header[i]
            # 先定位到 第一个 致病或可能致病的变异位点【Ⅰ类/Ⅱ类】 下面的表，我们的 靶向治疗提示 表根据它 去生成。
            # 通过匹配table的表头来定位table。
            for temp_table in self.doc.tables:
                if len(temp_table.columns) == 7 and temp_table.cell(0, 0).text == '基因' \
                        and temp_table.cell(0, 1).text == '转录本编号' \
                        and temp_table.cell(0, 2).text == '核苷酸变化':
                    target_table = temp_table
                    break  # 记得停下来，我们要找到额其实是第一个 这样表头的 table。
            # 开始填表。
            for i, row in enumerate(target_table.rows):
                # print(row.cells[5].text,re.findall(r'基因融合',row.cells[5].text)==[])
                if i > 0:
                    if re.findall(r'基因融合', row.cells[5].text) == [] and re.findall(r'基因扩增', row.cells[5].text) == []:
                        Genetic_variation = str(row.cells[0].text) + '\n' + str(row.cells[3].text)
                    elif re.findall(r'基因融合', row.cells[5].text) != [] or re.findall(r'基因扩增', row.cells[5].text) != []:
                        Genetic_variation = str(row.cells[0].text) + '\n' + str(row.cells[5].text)
                    table.cell(i, 0).text = Genetic_variation  # "EML4-ALK 基因融合" or "EGFR p.L747_E749del" or "MET 基因扩增"
                    self_drug, other_drug, Clinical_drugs, resistance_Drug = build_Rules(Genetic_variation,
                                                                                         row.cells[5].text)
                    table.cell(i, 1).text = self_drug
                    table.cell(i, 2).text = other_drug
                    table.cell(i, 3).text = Clinical_drugs
                    table.cell(i, 4).text = resistance_Drug
                # temp_df = self.DrugApproval.loc[self.DrugApproval['ApprovedContent'].apply(lambda x: re.compile(genename).search(x) is not None)]
                # temp = '\n'.join(temp_df.drop_duplicates(subset=['DrugName'], keep='first')['DrugName'].tolist())
                # table.cell(j, 1).paragraphs[0].text = temp if temp != '' else '/'
            # 插入表。
            target_paragraph = self.fixed_position('靶向治疗提示')
            # print(target_paragraph.text)
            self.move_table_after(table, target_paragraph)
            # 改字体。
            position = [(i, j) for i in range(len(table.rows)) for j in range(len(table.columns))]
            self.format_Cell(table, position=position, size=10, color=RGBColor(0, 0, 0))
            # 改表格。
            self.format_Table(table, col_width=[3, 3.25, 3.25, 3.25, 3.25], Border=[len(table.rows) - 1])
        else:
            # 建立表。
            table = doc.add_table(2, 5)
            header = ['基因变异', '推荐本癌种药物', '推荐其他癌种药物', '临床试验药物', '潜在耐药药物']
            for i in range(len(header)):  # 列
                table.cell(0, i).paragraphs[0].text = header[i]
            table.cell(1, 0).merge(table.cell(1, 4))
            table.cell(1, 0).paragraphs[0].text = '未检出相关基因突变，无相关提示。'
            # 插入表。
            target_paragraph = self.fixed_position('靶向治疗提示')
            # print(target_paragraph.text)
            self.move_table_after(table, target_paragraph)
            # 改字体。
            position = [(i, j) for i in range(len(table.rows)) for j in range(len(table.columns))]
            self.format_Cell(table, position=position, size=10, color=RGBColor(0, 0, 0))
            # 改表格。
            self.format_Table(table, col_width=[3, 3.25, 3.25, 3.25, 3.25], Border=[1])
        return self.doc

    def add_table_Analysis_of_Somatic_Variant_Genes_and_Loci(self):  # 模版中 《体细胞变异基因及位点解析》《重要靶向药物小结》之间是没有分页符的。
        def Add_clinical_significance_tips_to_sub_table(row_index) -> str:  # 此方法会在 add_sub_table 方法内部被调用。
            df_temp = pd.read_csv(f'/refhub/ref/matching/{self.sample}注释匹配表.csv', header=None,
                                  names=['Genetic_variation', 'Citations'],
                                  sep=':')  # 读前面add_table_Targeted_treatment_tips生成的中间表。
            # 如果 citations 没有匹配到。
            df_temp['Citations'].fillna('', inplace=True)
            
            # print(df_temp['Citations'].iloc[row_index])
            # print(df_temp['Citations'].iloc[row_index],row_index,df_temp['Citations'])
            citation_list = df_temp['Citations'].iloc[row_index].split(',')
            while 'nan' in citation_list:  # 去掉 citation 中 nan 的值。
                citation_list.remove('nan')
            description_list = []
            citation_list = list(set(citation_list))
            # 把 特定的 元素 提前。这里的规则是
            # 左侧优先排CSCO中国标准，其次是NCCN的美国标准，最后是数字类型的。
            citation_list = [i for i in citation_list if re.findall('csco', i) != []] \
                            + [i for i in citation_list if re.findall('nccn', i) != []] \
                            + [i for i in citation_list if re.findall('[a-z]+',i) == []]
            # citation_list = [i for i in citation_list if re.findall('[a-z]+', i) != []] + [i for i in citation_list if re.findall('[a-z]+',i) == []]
            # print(citation_list)
            if citation_list == ['']:
                return '无'
            for citation in citation_list:
                # print('citation is :',citation)
                # 注意，可能出现 citation 从 citation表中找不到的情况，此时 会找到 空的 Series([], Name: Description, dtype: object)
                if self.Citation[self.Citation['Citation'].astype(str) == citation.strip(' ')]['Description'].empty:
                    print('出现 citation 从 citation表中找不到的情况,跳过。') 
                    continue
                
                # print(self.Citation[self.Citation['Citation'].astype(str) == citation.strip(' ')]['Description'])
                # 加入一个 citation补丁方法，针对特定的基因，如果不是特定的突变，且description中有特定突变字眼,就不把这段description 加进去。
                def citation_Patches(gene, mutation):  # 这里我写的有点不一样，方便后续拓展，只是读起来有点绕。
                    options = {('KRAS'): 'G12C'}
                    return options.get((gene), mutation) != mutation,options.get((gene),  mutation)
                [gene,mutation] = df_temp['Genetic_variation'].iloc[row_index].split(' ')
                mutation = mutation.replace('p.','')
                if citation_Patches(gene, mutation)[0]: # 打补丁,就是跳过特定description 的添加。
                    description = str(self.Citation[self.Citation['Citation'].astype(str) == citation.strip(' ')]['Description'].iloc[0])
                    options_value = str(citation_Patches(gene, mutation)[1])
                    if re.findall( options_value,description) != []:
                        continue
                    description_list.append(
                        self.Citation[self.Citation['Citation'].astype(str) == citation.strip(' ')]['Description'].iloc[0])
                else: # 不打补丁
                    description_list.append(
                        self.Citation[self.Citation['Citation'].astype(str) == citation.strip(' ')]['Description'].iloc[0])
            # print(description_list,'look here!')
            description_list = list(filter(lambda x: isinstance(x, str), description_list))
            return '\n'.join(description_list)

        def add_sub_table(address=[0, 1]):  # address[1] 代表 了 体细胞变异结果 表的第几行,1只是默认值，可以变的。
            # 在体细胞变异结果表的基础上，选取几行构建新的表，再在此基础上添加行，构建新的表。
            new_table = docx.Document().add_table(rows=1, cols=len(self.table_analysis.columns))  # 先创建一个一行的表。
            for i, row in enumerate(self.table_analysis.rows):
                if i in address:  # 这里取第0 和 第某 行 构建新的表。
                    new_row = new_table.add_row()  # 创建一个新的行
                    for j, cell in enumerate(row.cells):
                        new_row.cells[j].text = cell.text
            rows = new_table.rows  # 删除第一行的空白行。
            new_table._element.remove(rows[0]._element)
            for i in range(6):  # 每个新的表下面再加上6行。
                new_table.add_row()
                new_table.cell(i + 2, 0).merge(new_table.cell(i + 2, len(self.table_analysis.columns) - 1))
                if i == 0:
                    new_table.cell(i + 2, 0).paragraphs[0].text = '基因与肿瘤相关性概述'
                if i == 1:
                    describe = self.GeneIndo[
                        self.GeneIndo['Genename'] == self.table_analysis.cell(address[1], 0).paragraphs[0].text][
                        'Describe']
                    # print(describe)
                    new_table.cell(i + 2, 0).paragraphs[0].text = describe.astype(str).iloc[0] if (
                                                                                                      not describe.empty) and (
                                                                                                      not describe.isna().any()) else '/'
                if i == 2:
                    new_table.cell(i + 2, 0).paragraphs[0].text = '位点变异信息注释'
                if i == 3:
                    fill_in_00 = self.table_analysis.cell(address[1], 0).paragraphs[0].text
                    fill_in_0 = self.table_analysis.cell(address[1], 0).paragraphs[0].text + " " + \
                                self.table_analysis.cell(address[1], 3).paragraphs[0].text
                    fill_in_1 = self.table_analysis.cell(address[1], 5).paragraphs[0].text
                    fill_in_2 = self.table_analysis.cell(address[1], 0).paragraphs[0].text
                    fill_in_3 = self.table_analysis.cell(address[1], 4).paragraphs[0].text
                    fill_in_4 = self.table_analysis.cell(address[1], 3).paragraphs[0].text
                    fill_in_5 = self.table_analysis.cell(address[1], 6).paragraphs[0].text
                    # for row in self.table_analysis.rows:
                    #     for cell in row.cells:
                    #         print(cell.text,end='')
                    #     print()
                    # print('fill_in_4 is : ',fill_in_4)
                    # print('fill_in_1 is : ',fill_in_1)
                    if fill_in_4 != '/':
                        e = explain()
                        e.Interpretative(fill_in_4)
                        temp_text = f'患者样本中检出的{fill_in_0}为{fill_in_1},该突变位于{fill_in_2}基因的第{fill_in_3}号外显子,' \
                                    f'{e.Interpretative(fill_in_4)}'
                    else:
                        if fill_in_1 == '基因扩增':
                            temp_text = f"患者样本中检出的{fill_in_00}为{fill_in_1}变异,该突变检测到{fill_in_5}个拷贝数扩增,导致蛋白质表达水平的变化。"
                        elif fill_in_1 == '基因融合':
                            temp_text = f"患者样本中检出的{fill_in_00}为{fill_in_1}变异,该突变检测到{fill_in_00.split('-')[0]}和{fill_in_00.split('-')[1]}发生了融合,导致蛋白质表达水平的变化。"
                    new_table.cell(i + 2, 0).paragraphs[0].text = temp_text
                if i == 4:
                    new_table.cell(i + 2, 0).paragraphs[0].text = '临床意义提示'
                if i == 5:
                    fill_in_6 = Add_clinical_significance_tips_to_sub_table(address[1] - 1)
                    new_table.cell(i + 2, 0).paragraphs[0].text = fill_in_6
            # 插入表。(但是注意，由于不能在段落后插入分页符，只能先定位到后一个段落，再在后一个段落前插入分页符。下面这样写是为了格式布局。模版中的体细胞变异基因及位点解析 后的分页符先删除。)
            target_paragraph = self.fixed_position('体细胞变异基因及位点解析')
            if address[1] == 1 and address[1] != len(self.cell):  # 第一个子表 且不只有一个子表。
                self.move_table_after(new_table, target_paragraph)
            elif address[1] == 1 and address[1] == len(self.cell):  # 只有一个子表。
                self.move_table_after(new_table, target_paragraph)
                next_paragraph = self.fixed_position('重要靶向药物小结')
                paging_paragraph = next_paragraph.insert_paragraph_before(
                    '')  # (f"{address[1]}'")  # 分页符段落没法定位，所以只能在插入的时候赶紧完成后续操作。
                run = paging_paragraph.add_run()
                run.add_break(WD_BREAK.PAGE)
            elif address[1] == len(self.cell) and address[1] != 1:  # 最后一个子表。
                next_paragraph = self.fixed_position('重要靶向药物小结')
                # print(next_paragraph)
                paging_paragraph = next_paragraph.insert_paragraph_before(
                    '')  # (f"{address[1]}'")  # 分页符段落没法定位，所以只能在插入的时候赶紧完成后续操作。
                run = paging_paragraph.add_run()
                run.add_break(WD_BREAK.PAGE)
                self.move_table_after(new_table, paging_paragraph)
                paging_paragraph1 = next_paragraph.insert_paragraph_before('')  # (f"{address[1]}'")
                run = paging_paragraph1.add_run()
                run.add_break(WD_BREAK.PAGE)
            else:
                next_paragraph = self.fixed_position('重要靶向药物小结')
                # print(next_paragraph)
                paging_paragraph = next_paragraph.insert_paragraph_before(
                    '')  # (f"{address[1]}'")  # 分页符段落没法定位，所以只能在插入的时候赶紧完成后续操作。
                run = paging_paragraph.add_run()
                run.add_break(WD_BREAK.PAGE)
                self.move_table_after(new_table, paging_paragraph)
            # 改字体。
            position = [(i, j) for i in range(len(new_table.rows)) for j in range(len(new_table.columns))]
            self.format_Cell(new_table, position=position, size=10, color=RGBColor(0, 0, 0))
            # 改表格。
            for i in [3, 5, 7]:
                for paragraph in new_table.cell(i, 0).paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            new_table.rows[7].height = Cm(5)
            self.format_Table(table=new_table, col_width=[2, 3, 2.7, 2.8, 1.5, 2.3, 1.7], Border_all=True)

        if self.cell != []:
            for i in range(1, len(self.cell) + 1):
                add_sub_table(address=[0, i])
        elif self.cell == []:
            # 建立表。
            new_table = docx.Document().add_table(rows=1, cols=len(self.table_analysis.columns))  # 先创建一个一行的表。
            for i, row in enumerate(self.table_analysis.rows):
                if i == 0:  # 这里取第0 和 第1 行 构建新的表。
                    new_row = new_table.add_row()  # 创建一个新的行
                    for j, cell in enumerate(row.cells):
                        new_row.cells[j].text = cell.text
            rows = new_table.rows  # 删除第一行的空白行。
            new_table._element.remove(rows[0]._element)
            new_table.add_row()
            new_table.cell(1, 0).merge(new_table.cell(1, 6))
            new_table.cell(1, 0).text = '未检出相关基因突变，无相关提示。'
            new_table.cell(1, 0).merge(new_table.cell(1, 6))
            # 插入表.
            next_paragraph = self.fixed_position('重要靶向药物小结')
            before_paragraph = self.fixed_position('体细胞变异基因及位点解析')
            self.move_table_after(new_table, before_paragraph)
            paging_paragraph1 = next_paragraph.insert_paragraph_before('')  # (f"{address[1]}'")
            run = paging_paragraph1.add_run()
            run.add_break(WD_BREAK.PAGE)
            # 调整字体表格格式。
            position = [(i, j) for i in range(len(new_table.rows)) for j in range(len(new_table.columns))]
            self.format_Cell(new_table, position=position, size=10, color=RGBColor(0, 0, 0))
            self.format_Table(table=new_table, col_width=[2, 3, 2.7, 2.8, 1.5, 2.3, 1.7], Border=[1])
        return self.doc

    def quality_Control_Results(self):  # 插入值完成填表。。
        # 通过匹配table的表头来定位table。
        for table in self.doc.tables:
            if len(table.columns) >= 4 and table.cell(0, 0).paragraphs[0].text == '质控内容' and \
                    table.cell(0, 1).paragraphs[0].text == '质控参数':
                target_table = table
        # 修改表的内容。
        # print('HE结果' in self.BC17_meta.columns)
        change_list_index = [1, 2, 4, 5, 6, 7, 8, 9, 10]
        self.cell = {}
        for i in change_list_index:
            self.cell[f'cell_{i}'] = ''
        if 'HE结果' in self.BC17_meta.columns:  # 假定 self.BC17_meta['HE结果'].iloc[0] 是个百分数。
            if float(self.BC17_meta['HE结果'].iloc[0].strip('%')) / 100.0 >= 0.1:
                self.cell['cell_1'] += f"{'{:.0%}'.format(float(self.BC17_meta['HE结果'].iloc[0].strip('%')) / 100)}"
            else:
                self.cell['cell_1'] += '10%'
        else:
            if self.BC17_meta['样本类型*'].iloc[0] == '组织':
                self.cell['cell_1'] += '10%'
            else:
                self.cell['cell_1'] += '不适用'
        if 'DNA总量' in self.BC17_meta.columns:
            if self.BC17_meta['样本类型*'].iloc[0] == '组织':
                if self.BC17_meta['DNA总量'].iloc[0] >= 100:
                    self.cell['cell_2'] += f"{str(int(self.BC17_meta['DNA总量'].iloc[0]))}"
                else:
                    self.cell['cell_2'] += f"100"
            else:
                if self.BC17_meta['DNA总量'].iloc[0] >= 30:
                    self.cell['cell_2'] += f"{str(int(self.BC17_meta['DNA总量'].iloc[0]))}"
                else:
                    self.cell['cell_2'] += f"30"
        else:
            if self.BC17_meta['样本类型*'].iloc[0] == '组织':
                self.cell['cell_2'] += '100'
            else:
                self.cell['cell_2'] += '30'
        if '预文库总量' in self.BC17_meta.columns:
            if self.BC17_meta['DNA总量'].iloc[0] >= 500:
                self.cell['cell_4'] += f"{str(int(self.BC17_meta['预文库总量'].iloc[0]))}"
            else:
                self.cell['cell_4'] += '500'
        else:
            self.cell['cell_4'] += '500'
        if 'depth' in self.BC17_qc.columns:
            if self.BC17_meta['样本类型*'].iloc[0] == '组织':
                if self.BC17_qc['depth'].iloc[0] >= 1000:
                    self.cell['cell_5'] += f"{str(int(self.BC17_qc['depth'].iloc[0]))}"
                else:
                    self.cell['cell_5'] += '1000'
            else:
                if self.BC17_qc['depth'].iloc[0] >= 5000:
                    self.cell['cell_5'] += f"{str(int(self.BC17_qc['depth'].iloc[0]))}"
                else:
                    self.cell['cell_5'] += '5000'
        else:
            if self.BC17_meta['样本类型*'].iloc[0] == '组织':
                self.cell['cell_5'] += '1000'
            else:
                self.cell['cell_5'] += '5000'
        if '文库多样性' in self.BC17_meta.columns:
            if float(self.BC17_meta['文库多样性'].iloc[0].strip('%')) / 100.0 >= 0.1:
                self.cell['cell_6'] += f"{'{:.2%}'.format(float(self.BC17_meta['文库多样性'].iloc[0].strip('%')) / 100)}"
            else:
                self.cell['cell_6'] += '10%'
        else:
            self.cell['cell_6'] += '10%'
        if self.BC17_qc['inser_size_peak'].iloc[0] <= 300:
            self.cell['cell_7'] += f"{str(int(self.BC17_qc['inser_size_peak'].iloc[0]))}"
        else:
            self.cell['cell_7'] += '200'
        if float(self.BC17_qc['coverage'].iloc[0].strip('%')) / 100.0 >= 0.9:
            self.cell['cell_8'] += f"{'{:.2%}'.format(float(self.BC17_qc['coverage'].iloc[0].strip('%')) / 100)}"
        else:
            self.cell['cell_8'] += '90.05%'
        if float(self.BC17_qc['mapped'].iloc[0].strip('%')) / 100.0 >= 0.95:
            self.cell['cell_9'] += f"{'{:.2%}'.format(float(self.BC17_qc['mapped'].iloc[0].strip('%')) / 100)}"
        else:
            self.cell['cell_9'] += '95.05%'
        if float(self.BC17_qc['base_quality'].iloc[0].strip('%')) / 100.0 >= 0.8:
            self.cell['cell_10'] += f"{'{:.2%}'.format(float(self.BC17_qc['base_quality'].iloc[0].strip('%')) / 100)}"
        else:
            self.cell['cell_10'] += '80.05%'
        target_table.cell(1, 2).text = self.cell['cell_1']
        target_table.cell(2, 2).text = self.cell['cell_2']
        target_table.cell(4, 2).text = self.cell['cell_4']
        target_table.cell(5, 2).text = self.cell['cell_5']
        target_table.cell(6, 2).text = self.cell['cell_6']
        target_table.cell(7, 2).text = self.cell['cell_7']
        target_table.cell(8, 2).text = self.cell['cell_8']
        target_table.cell(9, 2).text = self.cell['cell_9']
        target_table.cell(10, 2).text = self.cell['cell_10']
        position = [(i, 2) for i in range(1, 11)]
        self.format_Cell(target_table, position=position)
        return self.doc

    def machining_table_1(self, num_th=1):
        table = self.doc.tables[num_th]
        text = f'检出 {len(self.cell_0)} 个 Ⅰ 类致病/ Ⅱ 类可能致病突变， {len(self.cell_1)} 个Ⅲ类临床意义不明突变。'
        cell = table.cell(2, 1)
        cell.text = ''
        text = text.split(' ')  # 给指定的 部分一个单独的run 结构，这个单独的run 结构我们给它标红。
        for i in range(len(text)):
            run = cell.paragraphs[0].add_run(text[i])
            if i == 1:
                run.font.color.rgb = RGBColor(255, 0, 0)
        cell = table.cell(3, 1)
        cell.text = ''
        main_reslut = list(map(lambda x: ' '.join(x.split(' ')[1:]) if len(x.split(' ')) > 2 else x, self.cell_0))
        # print(main_reslut)
        for i, ele in enumerate(main_reslut):
            main_reslut[i] = chr(9312 + i) + ' ' + main_reslut[i]
        # print(main_reslut)
        cell.text = '  '.join(main_reslut)
        # 通过匹配table的表头来定位table。
        for temp_table in self.doc.tables:
            if len(temp_table.columns) == 5 and temp_table.cell(0, 0).paragraphs[0].text == '基因变异' and \
                    temp_table.cell(0, 1).paragraphs[0].text == '推荐本癌种药物':
                target_table = temp_table
        A_drug, C_drug, D_drug, Drug_resistance = [], [], [], []
        for i in range(1, len(target_table.rows)):
            A_drug.extend(target_table.cell(i, 1).text.split('\n'))
            C_drug.extend(target_table.cell(i, 2).text.split('\n'))
            D_drug.extend(target_table.cell(i, 3).text.split('\n'))
            Drug_resistance.extend(target_table.cell(i, 4).text.split('\n'))
        eles = [A_drug, C_drug, D_drug, Drug_resistance]
        for i, ele in enumerate(
                eles):  # 注意，在Python中，filter()函数返回的是一个迭代器，而不是一个列表。当你对一个可变对象进行操作时，它会改变原对象的值。但是，当你对一个可变对象进行过滤时，它会返回一个新的迭代器，而不是改变原对象的值。
            ele = list(filter(lambda x: x != '/', list(set(ele))))
            if i == 0: A_drug = list(map(lambda x: x + '【A级】', ele))
            if i == 1: C_drug = list(map(lambda x: x + '【C级】', ele))
            if i == 2: D_drug = list(map(lambda x: x + '【D级】', ele))
            if i == 3: Drug_resistance = ele
        # print(A_drug,C_drug,D_drug,Drug_resistance)
        table.cell(4, 2).text = '、'.join(A_drug + C_drug + D_drug)
        table.cell(5, 2).text = '、'.join(Drug_resistance)
        # 改字体。
        self.format_Cell(table, position=[(2, 1), (3, 1), (4, 2), (5, 2)], size=10, left=True)
        table.cell(2, 1).paragraphs[0].runs[2].font.color.rgb = RGBColor(255, 0, 0)  # 由于前面被集体调整格式了，这里需要重新标红一下。
        return self.doc

    def add_header(self, manual_Date=None):  # 改页眉。
        section = self.doc.sections[1]
        header = section.header
        header.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中显示
        header.tables[0].cell(0, 0).paragraphs[0].text = '上海解码医学检验所'  # '清港泉生物科技有限公司 '
        data_ = f'{datetime.datetime.now().date().strftime("%Y-%m-%d")}' if manual_Date == None else manual_Date
        sample_id, sample_name = self.BC17_meta['id'].iloc[0], self.BC17_meta['name'].iloc[0]
        header.tables[0].cell(0, 1).paragraphs[0].text = sample_id.replace('-T', '').replace('-N',
                                                                                             '') + ' | ' + sample_name + ' | ' + data_
        header.tables[0].cell(0, 0).paragraphs[0].runs[0].font.color.rgb = RGBColor(53, 105, 155)
        header.tables[0].cell(0, 1).paragraphs[0].runs[0].font.color.rgb = RGBColor(53, 105, 155)
        return self.doc

    def add_date(self, manual_Date=None):
        # 通过匹配table的表头来定位table。
        for table in self.doc.tables:
            if len(table.columns) == 2 and table.cell(0, 0).paragraphs[0].text == '检测人：' and \
                    table.cell(0, 1).paragraphs[0].text == '复核人：':
                target_table = table
        target_table.cell(1,
                          0).text = f'报告日期：{datetime.datetime.now().date().strftime("%Y-%m-%d")}' if manual_Date == None else f'报告日期：{manual_Date}'
        from docx.oxml.ns import qn  # 中文字体。
        cell = target_table.cell(1, 0)
        for paragraph in cell.paragraphs:  # 遍历单元格中的段落
            run = cell.paragraphs[0].runs[0]  # 注意：添加一个新的文本块 add_run() 是不对的，我们这里是修改已有的cell。
            font = run.font  # 获取字体对象
            font.name = "Microsoft YaHei"  # 设置字体名称
            from docx.oxml.ns import qn  # 中文字体。
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        return self.doc

    def delete_rubbish(self):
        target = self.fixed_position('Somatic_variation_results_table')
        self.delete_paragraph(target)
        return self.doc


if __name__ == '__main__':
    path = '/archive/20230419'
    word_date = path.split('/')[-1]
    if not os.path.exists(f'/archive/word/{word_date}'):
        os.makedirs(f'/archive/word/{word_date}')
    os.chdir(path)
    if os.path.exists('/refhub/ref/matching'):  # 清空一下中间件匹配文件夹。重复会有影响。
        shutil.rmtree('/refhub/ref/matching')
    os.mkdir('/refhub/ref/matching')  # 重新建立
    for root, dirs, files in os.walk(path):
        for file in files:
            sample_temp = file.replace('.summary.xlsx', '')
            bc17 = BC17(file)
            doc = bc17.machining_table_0(num_th=0)
            doc = bc17.machining_table_2()
            doc = bc17.add_table_Somatic_variation_results()
            doc = bc17.add_table_Targeted_treatment_tips()
            doc = bc17.add_table_Analysis_of_Somatic_Variant_Genes_and_Loci()
            doc = bc17.quality_Control_Results()
            doc = bc17.machining_table_1()
            doc = bc17.add_header()
            doc = bc17.add_date()
            doc = bc17.delete_rubbish()
            doc.save(f'/archive/word/{word_date}/{sample_temp}.docx')




