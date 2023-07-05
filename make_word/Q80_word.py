#coding=utf8
import docx, copy, re, os,datetime,math,shutil,logging,inspect,sys
from typing import List
import pandas as pd
import numpy as np
from docx import Document
from functools import wraps
from docx.oxml import CT_P, CT_Tbl
from docx.table import _Cell, Table, _Row
from docx.text.paragraph import Paragraph
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm
from docx.shared import RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import warnings
warnings.filterwarnings('ignore')

# 常见报错，以及 处理
#1、 "NoneType object is not iterable " 很可能是position的参数传递没有 写完整。
#2、"IndexError：collumn index is out of range" 很可能是 self.doc.add_table  写成了 docx.Document.add_table。
pd.set_option('display.max_columns', None)

class explain():  # 解读 c.96+1G>T 以及 p.L747_P753delinsS 的意思。
    def __init__(self):
        self.kind, self.position, self.explain = '', '', ''
        self.Amino_acid_dict = {}
        Amino_acid_list = ['甘氨酸', '丙氨酸', '缬氨酸', '亮氨酸', '异亮氨酸', '脯氨酸', '苯丙氨酸', '酪氨酸', '色氨酸', '丝氨酸', '苏氨酸', '半胱氨酸', '蛋氨酸',
                           '天冬酰胺', '谷氨酰胺', '天冬氨酸', '谷氨酸', '赖氨酸', '精氨酸', '组氨酸']
        abbreviate_table = ['G', 'A', 'V', 'L', 'I', 'P', 'F', 'Y', 'W', 'S', 'T', 'C', 'M', 'N', 'Q', 'D', 'E', 'K',
                            'R', 'H']
        for i in range(len(Amino_acid_list)):
            self.Amino_acid_dict[abbreviate_table[i]] = Amino_acid_list[i]

    def classification_Description(self, raw_explain):
        result = []
        if re.match('delins', raw_explain) != None:
            result.append(f"缺失，同时发生了{len(re.match('delins([A-WY-Z]+)', raw_explain).group(1))}个氨基酸的插入")
            raw_explain = raw_explain.replace(re.match('(delins[A-WY-Z]+)', raw_explain).group(1), '')
        if re.match('del', raw_explain) != None:
            result.append('发生了缺失')
            raw_explain = raw_explain.replace(re.match('(del)', raw_explain).group(1), '')
        if re.match('ins', raw_explain) != None:
            result.append(f"发生了{len(re.match('ins([A-WY-Z]+)', raw_explain).group(1))}个氨基酸的插入")
            raw_explain = raw_explain.replace(re.match('(ins[A-WY-Z]+)', raw_explain).group(1), '')
        if re.match('[A-WY-Z]+', raw_explain) != None:
            result.append(f"突变为{self.Amino_acid_dict[re.match('([A-WY-Z]+)', raw_explain).group(1)]}。")
            raw_explain = raw_explain.replace(re.match('([A-WY-Z]+)', raw_explain).group(1), '')
        if re.match('[X]', raw_explain) != None:
            result.append(f"突变为 未知类型氨基酸。")
        if re.match('fs', raw_explain) != None:
            i = re.match(r'fs\*([\d]+)', raw_explain).group(1)
            result.append(f'出现移码变异,导致后续的氨基酸序列发生错位，从而改变了翻译后的蛋白质序列，并且在插入位置之后的{i}位形成了一个终止密码子'
                          f'导致蛋白质合成的过程中到达这个位置后会停止合成。这种类型的变异通常会导致蛋白质合成的过程提前终止，从而产生缺失的蛋白质片段。'
                          f'由于蛋白质的功能通常与其结构密切相关，这种缺失可能会对蛋白质的功能和稳定性产生影响。')
            raw_explain = raw_explain.replace(re.match(r'(fs\*[\d]+)', raw_explain).group(1), '')
        return ''.join(result)

    def Interpretative_p(self, sentence):
        match = re.match(r'(^[\w]+[0-9]+)', sentence)  # match.group(1)=L747_P753
        match_group_list = match.group(1).split('_')
        # print(match_group_list)
        if len(match_group_list) > 1:
            self.position = '第' + match_group_list[0][1:] + '位' + self.Amino_acid_dict[
                match_group_list[0][0]] + '至' + '第' + match_group_list[1][1:] + '位' + self.Amino_acid_dict[
                                match_group_list[1][0]]
            raw_explain = sentence.replace(match.group(1), '')  # delinsD
            self.explain = self.classification_Description(raw_explain)
        else:
            self.position = '第' + match_group_list[0][1:] + '位' + self.Amino_acid_dict[match_group_list[0][0]]
            raw_explain = sentence.replace(match.group(1), '')
            try:
                self.explain = self.classification_Description(raw_explain)
            except:
                self.explain = '没找到'
        return f'{self.kind}{self.position}{self.explain}'

    def Interpretative_c(self, sentence):
        return '导致该基因编码的蛋白序列未知，预测产生无功能蛋白，该基因正常蛋白功能丧失。'

    def Interpretative(self, sentence):
        if sentence and sentence[0] == 'p':
            self.kind = '导致该基因编码的蛋白序列的'
            # print(self.Interpretative_p(sentence[2:]))
            return self.Interpretative_p(sentence[2:])
        elif sentence and sentence[0] == 'c':
            self.kind = '导致该基因编码的蛋白序列未知'
            # print(self.Interpretative_c(sentence[2:]))
            return self.Interpretative_c(sentence[2:])
        else:
            self.kind = '导致该基因编码的蛋白序列未知'
            return '未知.'

class Tools():
    def Set_Background_Color(self, cell, rgbColor):
        shading_elm = parse_xml(
            r'<w:shd {} w:fill="{color_value}"/>'.format(nsdecls('w'), color_value=rgbColor))  # 换背景色固定写法，照抄即可
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def delete_paragraph(self, paragraph):  # 删除段落。
        p = paragraph._element
        p.getparent().remove(p)
        paragraph._p = paragraph._element = None

    # 在目标段落前面插入。
    def inserted_paragraph(self,doc,behind_paragraph_text,insert_text):
        target_paragraph = self.fixed_position(doc,behind_paragraph_text)
        table_before_paragraph_0 = target_paragraph.insert_paragraph_before(insert_text)
        return doc

    def move_table_after(self, doc, paragraph_end_with,table_index=-1):  # 从指定paragraph 后insert table。
        tables = doc.tables
        table = tables[table_index]  # 默认处理doc 中的最后一个table。
        paragraph = self.fixed_position(doc, paragraph_end_with)
        tbl, p = table._tbl, paragraph._p
        p.addnext(tbl)
        # 跨页获取表头。（这个很重要。）
        def set_repeat_table_header(row):
            tr = row._tr  # 获取行的底层XML元素
            trPr = tr.get_or_add_trPr()  # 获取或添加行属性
            tblHeader = OxmlElement('w:tblHeader')  # 创建w:tblHeader元素
            tblHeader.set(qn('w:val'), "true")  # 设置w:tblHeader元素的属性值为"true"
            trPr.append(tblHeader)  # 将w:tblHeader元素添加到行属性中
            return row
        set_repeat_table_header(table.rows[0])
        return doc

    def move_table_after_supplement(self,paragraph,table_index=-1):
        tables = doc.tables
        table = tables[table_index]  # 默认处理doc 中的最后一个table。
        tbl, p = table._tbl, paragraph._p
        p.addnext(tbl)

    def Set_Background_Color(self, cell, rgbColor):
        shading_elm = parse_xml(
            r'<w:shd {} w:fill="{color_value}"/>'.format(nsdecls('w'), color_value=rgbColor))  # 固定写法，照抄即可
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def fill_empty_cells_with_slash(self,doc,table_index=-1):
        tables = doc.tables
        table = tables[table_index]
        for row in table.rows:
            for cell in row.cells:
                if not cell.text.strip() or cell.text.strip() == '.':
                    cell.text = '/'
        return doc

    def fixed_position(self,doc, End_character):  # End_character 用类似于 'pathogenictable'。
        for paragraph in doc.paragraphs:
            paragraph_text = paragraph.text
            if paragraph_text.endswith(End_character):  # 定位目标段落。
                target = paragraph
                return target

    # 把一个cell中原本 \n 模拟分段的内容改为真正的分段，并指定段间距6磅；
    def format_Cell_Space(self,doc,table_index=-1,position=[]):  # doxc 中 table 和 table 中的cell 都是可变对象。
        tables = doc.tables
        table = tables[table_index]  # 默认处理doc 中的最后一个table。
        def set_paragraph_spacing(paragraph):
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_after = Pt(0)
            paragraph_format.line_spacing = Pt(10) # 设置行距6磅
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        # print('position is ',position)
        for (row, column) in position:
            cell = table.cell(row, column)
            paragraphs = cell.paragraphs
            print("段落数量：", len(paragraphs))
            for i, paragraph in enumerate(paragraphs):
                print(f"段落 {i + 1}: {paragraph.text}")
            text = cell.text
            # 清空单元格中的全部内容
            for paragraph in cell.paragraphs:
                cell._element.remove(paragraph._element)
            paragraphs = text.split('\n')
            for paragraph_text in paragraphs:
                paragraph = cell.add_paragraph()
                paragraph.text = paragraph_text
                set_paragraph_spacing(paragraph)
        return doc

    # 插值不用 设置 table_index 用默认的就好了，修改表的话，就要设置 table_index.
    # 设置除了 第 column_index 列 ，从第二行起非居中对齐。
    def format_Cell(self,doc,table_index=-1,position=[],size=10, color=RGBColor(0, 0, 0),left=False,no_center_num=None):  # doxc 中 table 和 table 中的cell 都是可变对象。
        tables = doc.tables
        table = tables[table_index]  # 默认处理doc 中的最后一个table。
        # print('position is ',position)
        for (row, column) in position:
            cell = table.cell(row, column)
            for paragraph in cell.paragraphs:  # 遍历单元格中的段落
                for run in cell.paragraphs[0].runs:
                    # run = cell.paragraphs[0].runs[0]  # 注意：添加一个新的文本块 add_run() 是不对的，我们这里是修改已有的cell。
                    font = run.font  # 获取字体对象
                    font.name = "Microsoft YaHei"  # 设置字体名称为宋体
                    font.size = Pt(size)  # 设置字体大小为12号
                    font.color.rgb = color  # 设置字体颜色为黑色
                    if left == False:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 设置段落水平居中
                    elif left == True:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    from docx.oxml.ns import qn  # 中文字体。
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # 设置单元格垂直居中
        # 设置除了 第 column_index 列 ，从第二行起非居中对齐。
        def set_column_alignment(table, column_index):
            for i, row in enumerate(table.rows):
                if i > 0:
                    cell = row.cells[column_index]
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        set_column_alignment(table, no_center_num) if no_center_num!=None else print('')
        return doc

    def format_Table(self, doc,table_index=-1, header_size=10,First_column_colore=True,col_width=[1, 1], header_color=RGBColor(0, 0, 0), Border=[],
                     Border_all=False,no_colore_column=[]):
        tables = doc.tables
        table = tables[table_index]  # 默认处理doc 中的最后一个table。
        # 遍历每一行
        for i, row in enumerate(table.rows):
            # 获取每一行的所有单元格
            cells = row.cells
            # 遍历每个单元格
            row.height = Cm(1)  # 统一行高
            for index,cell in enumerate(cells):
                # 获取单元格的段落对象
                paragraph = cell.paragraphs[0]
                # 如果 第二行是 “未检出相关基因突变”，设置单元格高度。
                if cell.paragraphs[0].text == '未检出相关基因突变' or cell.paragraphs[0].text == '未检出相关基因突变，无相关提示。':
                    table.rows[1].height = Cm(1)
                # 如果是首行，设置字体为粗体，并设置背景色为灰色 ，设置单元格高度。
                if i == 0:
                    run = paragraph.runs[0]
                    run.font.bold = True
                    run.font.size = Pt(header_size)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    table.rows[0].height = Cm(1)
                    self.Set_Background_Color(cell, "CFDAE6")
                # 如果是奇数行，设置背景色为浅蓝色,且第一列不变色。 First_column_colore 来控制第一列是否变色。
                elif i % 2 != 1 and First_column_colore==True:  # and cell !=cells[0]:
                    if no_colore_column==[]: # 没有 no_colore_column 参数时 就按照原逻辑来，有no_colore_column参数，优先考虑no_colore_column参数。
                        self.Set_Background_Color(cell, "F1F1F1")
                    else:
                        if index not in no_colore_column:
                            self.Set_Background_Color(cell, "F1F1F1")
                elif i % 2 != 1 and First_column_colore==False:  # and cell !=cells[0]:
                    if no_colore_column==[]: # 没有 no_colore_column 参数时 就按照原逻辑来，有no_colore_column参数，优先考虑no_colore_column参数。
                        if index != 0:
                            self.Set_Background_Color(cell, "F1F1F1")
                    else:
                        if index not in no_colore_column:
                            self.Set_Background_Color(cell, "F1F1F1")
                if cell.text == '基因与肿瘤相关性概述' or cell.text == '位点变异信息注释' or cell.text == '临床意义提示':
                    for paragraph in cell.paragraphs:
                        run = cell.paragraphs[0].runs[0]
                        font = run.font
                        font.bold = True
                        font.name = "Microsoft YaHei"  # 设置字体名称
                        font.size = Pt(11)  # 设置字体大小为12号
                        font.color.rgb = RGBColor(67, 114, 161)  # 设置字体颜色为黑色
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 设置段落水平左对齐。
                        from docx.oxml.ns import qn  # 中文字体。
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # 设置单元格垂直居中
        # # 按比例列宽调整。
        # for i in range(len(table.rows)):
        #     for col in range(len(table.columns)):
        #         table.cell(i, col).width = Cm(col_width[col])
        # 按照固定的宽度调整。以免表格超出页边的距离。
        table.autofit = False
        table.allow_autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.width = Cm(16)
        for i, col in enumerate(col_width):
            table.columns[i].width = Cm(col)
        # 加边框
        def set_cell_border(cell, **kwargs):
            # 这里的kwargs 其实就是一个dict，key 是 top、bottom 这些，value是 {} 一个嵌套的dict。
            # 不定数量的形参传入时常用这样的写法，形参传入时=符号充当了dict的构建，=左边是key右边是value。
            """
            Set cell`s border
            Usage:
            set_cell_border(
                cell,
                top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
                bottom={"sz": 12, "color": "#00FF00", "val": "single"},
                left={"sz": 24, "val": "dashed", "shadow": "true"},
                right={"sz": 12, "val": "dashed"},
            )
            """
            # XML 元素是 XML 文档中必不可缺的部分，我们可以将 XML 元素看成一个容器，其中存放了文本，元素，属性，媒体对象或所有的这些。XML 文档包含 XML 元素。
            tc = cell._tc  # cell._tc 是一个单元格对象。_tc 属性是一个内部属性，它返回一个 _Element 对象，表示单元格的 XML 元素。
            tcPr = tc.get_or_add_tcPr()  # 获取短元格内部属性。
            # check for tag existnace, if none found, then create one
            tcBorders = tcPr.first_child_found_in(
                "w:tcBorders")  # 它返回了单元格属性对象中第一个 <w:tcBorders> 元素。如果该元素不存在，则该方法返回 None。
            if tcBorders is None:
                tcBorders = OxmlElement('w:tcBorders')
                tcPr.append(tcBorders)
            # list over all available tags
            for edge in ('left', 'top', 'right', 'bottom', 'insideH', 'insideV'):
                edge_data = kwargs.get(edge)  # edge_data 就是形参中内部嵌套的子dict。
                if edge_data:
                    tag = 'w:{}'.format(edge)  # w:left
                    # check for tag existnace, if none found, then create one
                    from docx.oxml.ns import qn
                    element = tcBorders.find(qn(tag))  # qn(tag) 是一个函数调用，它将字符串 tag 转换为一个 XML 命名空间。
                    if element is None:
                        element = OxmlElement(tag)
                        tcBorders.append(element)
                    # looks like order of attributes is important
                    for key in ["sz", "val", "color", "space", "shadow"]:
                        if key in edge_data:
                            element.set(qn('w:{}'.format(key)), str(edge_data[key]))  # 在元素中 把 xml 和 形参输入的值 对应起来。
        if Border == []:
            Border = [len(table.rows) - 1] # 把默认改成这个。
        if Border != [] and Border_all == False:  # Border = [] list 中存放 行号，第几行的所有单元格统一加某种边框。
            # 给指定单元格加边框
            for j in Border:
                for i in range(len(table.columns)):
                    set_cell_border(
                        table.cell(j, i),
                        # top={"color": "#000000", "space": "0"},
                        bottom={"val": "single", "color": "#CFDAE6", "space": "1"},
                        # left={"color": "#000000", "space": "0"},
                        # right={"sz": 0.5, "val": "double", "color": "#000000", "space": "0"},
                        # insideH={"sz": 0.5, "val": "double", "color": "#000000", "space": "0"},
                        # end={"sz": 0.5, "val": "double", "color": "#000000", "space": "0"}
                    )
        elif Border_all == True:
            # 给全部单元格加边框。
            for row in table.rows:
                for cell in row.cells:
                    set_cell_border(
                        cell,
                        top={"val": "dotted", "color": "#CFDAE6", "space": "0"},
                        bottom={"val": "dotted", "color": "#CFDAE6", "space": "0"},
                        left={"val": "dotted", "color": "#CFDAE6", "space": "0"},
                        right={"val": "dotted", "color": "#CFDAE6", "space": "0"},
                        insideH={"val": "dotted", "color": "#CFDAE6", "space": "0"},
                        end={"val": "dotted", "color": "#CFDAE6", "space": "0"}
                    )
        # table.autofit = True
        return doc

    # 像 指定 table_index 一样，修改的 target_paragraph 需要指定好是哪一个。
    def format_paragraph(self,doc, paragraph_end_with, size=10, bold=False):
        target_paragraph = self.fixed_position(doc,paragraph_end_with)
        run = target_paragraph.runs[0]
        run.font.name = "Microsoft YaHei"  # 设置字体名称为宋体
        run.font.size = Pt(size)  # 设置字体大小为12号
        run.font.bold = bold
        from docx.oxml.ns import qn  # 中文字体。
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        return doc

    # def Merge_cells_by_first_column(self,doc,table_index=-1, column_index=0):
    #     tables = doc.tables
    #     table = tables[table_index]
    #     first_column = table.columns[column_index]
    #     previous_text = first_column.cells[column_index].text
    #     merge_start_index = 0
    #     for i in range(1, len(first_column.cells)):
    #         current_text = first_column.cells[i].text
    #         if current_text != previous_text:
    #             # 合并单元格，并将第一个单元格设置为保留的值
    #             table.cell(merge_start_index, column_index).merge(table.cell(i - 1, column_index))
    #             table.cell(merge_start_index, column_index).text = previous_text
    #             merge_start_index = i
    #         else:
    #             # 将后续相同值的单元格清空
    #             table.cell(i, column_index).text = ""
    #         previous_text = current_text
    #     # 执行最后一个合并操作
    #     table.cell(merge_start_index, column_index).merge(table.cell(len(first_column.cells) - 1, column_index))
    #     table.cell(merge_start_index, column_index).text = previous_text
    #     return doc

    def Merge_cells_by_first_column(self,doc,table_index=-1, column_index=0):
        tables = doc.tables
        table = tables[table_index]
        first_column = table.columns[column_index]
        if len(table.rows)>1:
            previous_text = first_column.cells[1].text
            merge_start_index = 1
            previous_left_text = table.cell(1, column_index - 1).text if column_index > 0 else None
            print('table有{}行'.format(len(table.rows)))
            for i in range(2, len(first_column.cells)):
                #print(f'第{i}行','merge_start_index 是',merge_start_index)
                #print('table.cell(merge_start_index, column_index-1).text 是 ',table.cell(merge_start_index, column_index-1).text)
                #print('previous_left_text 是 ',previous_left_text)
                current_text = first_column.cells[i].text
                current_left_text = table.cell(i, column_index - 1).text if column_index > 0 else None
                #print('current_text 是 ', current_text, 'previous_text', previous_text)
                if (current_text != previous_text or current_left_text != previous_left_text)\
                        and (current_left_text==None or table.cell(merge_start_index, column_index-1).text==previous_left_text):
                    # 合并单元格，并将第一个单元格设置为保留的值
                    table.cell(merge_start_index, column_index).merge(table.cell(i - 1, column_index))
                    table.cell(merge_start_index, column_index).text = previous_text
                    merge_start_index = i
                    #print('一共有{}行，从0开始处理到第{}行前的合并'.format(len(first_column.cells),i))
                elif current_text == previous_text and (current_left_text==None or current_left_text == previous_left_text):
                    # 将后续相同值的单元格清空
                    table.cell(i, column_index).text = ""
                    #print('第{}行补了空值'.format(i))
                previous_text = current_text
                previous_left_text = current_left_text
            # 执行最后一个合并操作
            if current_text == previous_text and (current_left_text==None or current_left_text == previous_left_text):
                table.cell(merge_start_index, column_index).merge(table.cell(len(first_column.cells) - 1, column_index))
                table.cell(merge_start_index, column_index).text = previous_text
        return doc

    # def Merge_cells_by_first_column(self,doc,table_index=-1, column_index=0):
    #     tables = doc.tables
    #     table = tables[table_index]
    #     first_column = table.columns[column_index]
    #     previous_text = first_column.cells[1].text
    #     merge_start_index = 1
    #     for i in range(2, len(first_column.cells)):
    #         current_text = first_column.cells[i].text
    #         if current_text != previous_text:
    #             # 合并单元格，并将第一个单元格设置为保留的值
    #             table.cell(merge_start_index, column_index).merge(table.cell(i - 1, column_index))
    #             table.cell(merge_start_index, column_index).text = previous_text
    #             merge_start_index = i
    #         else:
    #             # 将后续相同值的单元格清空
    #             table.cell(i, column_index).text = ""
    #         previous_text = current_text
    #     # 执行最后一个合并操作
    #     if current_text == previous_text:
    #         table.cell(merge_start_index, column_index).merge(table.cell(len(first_column.cells) - 1, column_index))
    #         table.cell(merge_start_index, column_index).text = previous_text
    #     return doc

    def Specify_cell_color_change(self,doc,search_text,color=RGBColor(255, 0, 0),table_index=-1):
        tables = doc.tables
        table = tables[table_index]
        for row in table.rows:
            for cell in row.cells:
                if search_text in cell.text:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if search_text in run.text:
                                font = run.font
                                font.color.rgb = color # RGBColor(255, 0, 0)
        return doc

    def Specify_cell_color_change_condition(self,doc,color_address,condition_addre=(0,0),condition="!='未检出'",color=RGBColor(255, 0, 0),table_index=-1):
        tables = doc.tables
        table = tables[table_index]
        condition_addre_text = table.cell(condition_addre[0],condition_addre[1]).text
        if eval('condition_addre_text'+condition):
            for (row, column) in color_address:
                cell = table.cell(row, column)
                for paragraph in cell.paragraphs:  # 遍历单元格中的段落
                    for run in cell.paragraphs[0].runs:
                        font = run.font
                        font.color.rgb = color
        return doc

    def test(self,result,num):
        print(num)
        return result


def log_to_file(folder, filename):
    def decorator(func):
        @wraps(func)
        def wrapper(*args, **kwargs):
            # 配置日志记录器
            logging.basicConfig(level=logging.DEBUG, filename=f"{folder}/{filename}", filemode='a')  # a 是追加，w重写。
            logger = logging.getLogger(func.__name__)
            # 输出方法名称和参数信息
            logger.debug(f"Running {func.__name__} with args: {args}, kwargs: {kwargs}")
            # 执行方法
            result = func(*args, **kwargs)
            # 输出方法执行结果
            logger.debug(f"{func.__name__} result: {result}")
            return result
        return wrapper
    return decorator

# 规范化 工具的调用; 一定要 在 tools 类中的方法有 result参数，才能  用这个装饰器 规范化调用，tools 中的test 就是一个例子。
def Tools_Decorator(tool: str, *tool_args, **tool_kwargs):
    frame = inspect.currentframe().f_back # 使用inspect模块获取当前帧的上一帧，即调用test方法的帧。目的是识别 position=position_1还是position_0 等等。
    code = inspect.getframeinfo(frame).code_context[0].strip()
    # match = re.search(r'\((.*)position=(.*?)[\)|,]', code)
    match = re.search(r'\(\s*(.*?)\s*position\s*=\s*(.*?)\s*[\)|,]', code)
    if match:
        param = match.group(2)
        print(param)  # position_0
    else:
        param = None
    def tools_decorator(func):
        @wraps(func)
        def wrapper(*args, **kwargs):
            # 创建Tools实例
            tools = Tools()
            # 获取工具方法
            tool_func = getattr(tools, tool, None)
            if tool_func is None:
                raise ValueError(f"Invalid tool: {tool}")

            result = func(*args, **kwargs)
            if param:      # 装饰器 给了 position 的参数的情况。
                # print('position' in kwargs)
                # print(kwargs['position'],'方法中用到了position')
                # print(kwargs['position_1'], '方法中用到了position_1')
                # 替换position或position_1参数 利用了python 可变对象特性,一定是变，而不是重新赋值;
                position = kwargs.get('position', None)
                position_0 = kwargs.get('position_0', None)
                position_1 = kwargs.get('position_1', None)
                param_mapping = {
                    'position': position,'position_0': position_0,'position_1': position_1
                }
                tool_kwargs['position'] = param_mapping.get(param, None)

            processed_data = tool_func(result, *tool_args, **tool_kwargs)  # 传递tool_args和tool_kwargs
            return processed_data
        return wrapper
    return tools_decorator


class Q80():
    def __init__(self, input_file):
        self.doc = Document('/refhub/ref/masterplate/解码80基因报告母版.2023v1的副本.docx')
        self.input_file = input_file
        self.sample = self.input_file.split('/')[-1].replace('.summary.xlsx', '')
        print(self.sample, 'input_file is : ', input_file)
        df = pd.read_excel(input_file, sheet_name=None)
        self.BC17_meta, self.BC17_fusion, self.BC17_cnv, self.BC17_qc,self.Q80_msi,self.Q80_chemo = df['meta'].fillna(''), df['fusion'], df['cnv'], df['qc'],df['msi'],df['chemo']
        self.Q80_msi.columns = ['Total_Number_of_Sites', 'Number_of_Somatic_Sites', 'percent','4','5','6','7','8','9','10']
        self.Q80_chemo = self.Q80_chemo.drop('review', axis=1)
        self.BC17_snpindel = pd.concat([df['somatic'], df['germline']])
        # 以下为 出报告用到的参考表；
        df = pd.read_excel('/refhub/ref/drug/DrugCombinedPlan.xlsx',sheet_name=None)
        self.DrugCombinedPlan = df['Sheet1'].dropna(subset=['癌种'],how='any')
        df = pd.read_excel('/refhub/ref/drug/DrugApproval.xlsx', sheet_name=None)
        self.DrugApproval = df['药物获批情况表'].dropna(subset=['ApprovedContent'], how='any')
        df = pd.read_excel('/refhub/ref/gene/GeneInfo.xlsx', sheet_name=None)
        self.GeneIndo = df['Genename_table'].dropna(subset=['Genename'], how='any')
        df = pd.read_excel("/refhub/ref/citation/Citation.xlsx", sheet_name=None)
        self.Citation = df['citation']
        # 以下为 此样本内容属性；
        date_string = self.BC17_meta['到样日期*'].iloc[0]  # 2023-03-01。在这里，我们首先使用strptime()函数将字符串转换为datetime对象，然后使用strftime()函数将datetime对象转换为字符串并指定所需的格式。
        new_date_string =  datetime.datetime.strptime(date_string, "%Y/%m/%d").strftime("%Y-%m-%d")
        self.sample = self.BC17_meta['id'].iloc[0]
        self.name = self.BC17_meta['name'].iloc[0]
        self.gender = self.BC17_meta['gender'].iloc[0]
        self.age = self.BC17_meta['age'].astype(str).iloc[0]
        self.cancer = self.BC17_meta['cancer'].iloc[0]
        self.clinname = self.BC17_meta['clinname'].iloc[0]
        self.hospital = self.BC17_meta['送检医院'].iloc[0]
        self.arrival_date = new_date_string
        self.sample_type = self.BC17_meta['样本类型*'].iloc[0]
        self.panel = self.BC17_meta['panel'].iloc[0]
        self.projectname = self.BC17_meta['projectname'].iloc[0]
        self.report_template = self.BC17_meta['报告模版'].iloc[0]
        self.class_Tools = Tools()
        self.cell = []               # 内部存储 多个 Mutation_ele 类对象。
        df_ref_snpindel = self.BC17_snpindel[self.BC17_snpindel['review'].isin(['S1', 'S2'])]
        df_ref_fusion = self.BC17_fusion[self.BC17_fusion['review'].isin(['S1', 'S2'])]
        df_ref_cnv = self.BC17_cnv[self.BC17_cnv['review'].isin(['S1', 'S2'])]
        for index,row in df_ref_snpindel.iterrows():
            ele = self.Mutation_ele(gene=row['Gene.smallrefGene'],rna=row['RNA'],nucleotide=row['NCchange']\
                                    ,amino_acid=row['AAchange'],exon=row['EXON'],variant_type=row['ExonicFunc.smallrefGene'],vaf_cnv=row['VAF'])
            self.cell.append(ele)
        for index, row in df_ref_fusion.iterrows():
            ele = self.Mutation_ele(gene=row['Region1']+'-'+row['Region2'],rna='',nucleotide=''\
                                    ,amino_acid='',exon='',variant_type='基因融合',vaf_cnv='{:.2f}%'.format((float(row['Break_support1']) + float(row['Break_support2'])) / float(row['Total_depth']) * 100))
            self.cell.append(ele)
        for index, row in df_ref_cnv.iterrows():
            ele = self.Mutation_ele(gene=row['Gene'],rna=self.GeneIndo[self.GeneIndo['Genename'] == row['Gene']]['RNA'].iloc[0].split('.')[0]\
                                    ,nucleotide='',amino_acid='',exon=''\
                                    ,variant_type='基因扩增',vaf_cnv=row['CNN'])
            self.cell.append(ele)
        print('S1 S2 所有出问题的地方如下 ############(其中一部分可能会被过滤掉。)')
        for i in self.cell:
            print(i.gene,i.rna,i.nucleotide,i.amino_acid,i.exon,i.variant_type,i.vaf)
        print('S1 S2 所有出问题的地方如上 ############(其中一部分可能会被过滤掉。)')
        self.cell_g1g2 = []  # 内部存储 多个 Mutation_ele 类对象。
        df_ref_snpindel = self.BC17_snpindel[self.BC17_snpindel['review'].isin(['G1', 'G2'])]
        df_ref_fusion = self.BC17_fusion[self.BC17_fusion['review'].isin(['G1', 'G2'])]
        df_ref_cnv = self.BC17_cnv[self.BC17_cnv['review'].isin(['G1', 'G2'])]
        for index,row in df_ref_snpindel.iterrows():
            ele = self.Mutation_ele(gene=row['Gene.smallrefGene'],rna=row['RNA'],nucleotide=row['NCchange']\
                                    ,amino_acid=row['AAchange'],exon=row['EXON'],variant_type=row['ExonicFunc.smallrefGene'],vaf_cnv=row['VAF'])
            self.cell_g1g2.append(ele)
        for index, row in df_ref_fusion.iterrows():
            ele = self.Mutation_ele(gene=row['Region1']+'-'+row['Region2'],rna='',nucleotide=''\
                                    ,amino_acid='',exon='',variant_type='基因融合',vaf_cnv='{:.2f}%'.format((float(row['Break_support1']) + float(row['Break_support2'])) / float(row['Total_depth']) * 100))
            self.cell_g1g2.append(ele)
        for index, row in df_ref_cnv.iterrows():
            ele = self.Mutation_ele(gene=row['Gene'],rna=self.GeneIndo[self.GeneIndo['Genename'] == row['Gene']]['RNA'].iloc[0].split('.')[0]\
                                    ,nucleotide='',amino_acid='',exon=''\
                                    ,variant_type='基因扩增',vaf_cnv=row['CNN'])
            self.cell_g1g2.append(ele)
        print('G1 G2 所有出问题的地方如下 ############')
        for i in self.cell_g1g2:
            print(i.gene,i.rna,i.nucleotide,i.amino_acid,i.exon,i.variant_type,i.vaf)
        print('G1 G2 所有出问题的地方如上 ############')
        self.cell_S3S4g3g4 = []  # 内部存储 多个 Mutation_ele 类对象。
        df_ref_snpindel = self.BC17_snpindel[self.BC17_snpindel['review'].isin(['G1', 'G2'])]
        df_ref_fusion = self.BC17_fusion[self.BC17_fusion['review'].isin(['G1', 'G2'])]
        df_ref_cnv = self.BC17_cnv[self.BC17_cnv['review'].isin(['G1', 'G2'])]
        for index,row in df_ref_snpindel.iterrows():
            ele = self.Mutation_ele(gene=row['Gene.smallrefGene'],rna=row['RNA'],nucleotide=row['NCchange']\
                                    ,amino_acid=row['AAchange'],exon=row['EXON'],variant_type=row['ExonicFunc.smallrefGene'],vaf_cnv=row['VAF'])
            self.cell_S3S4g3g4.append(ele)
        for index, row in df_ref_fusion.iterrows():
            ele = self.Mutation_ele(gene=row['Region1']+'-'+row['Region2'],rna='',nucleotide=''\
                                    ,amino_acid='',exon='',variant_type='基因融合',vaf_cnv='{:.2f}%'.format((float(row['Break_support1']) + float(row['Break_support2'])) / float(row['Total_depth']) * 100))
            self.cell_S3S4g3g4.append(ele)
        for index, row in df_ref_cnv.iterrows():
            ele = self.Mutation_ele(gene=row['Gene'],rna=self.GeneIndo[self.GeneIndo['Genename'] == row['Gene']]['RNA'].iloc[0].split('.')[0]\
                                    ,nucleotide='',amino_acid='',exon=''\
                                    ,variant_type='基因扩增',vaf_cnv=row['CNN'])
            self.cell_S3S4g3g4.append(ele)
        print('S3S4g3g4 所有出问题的地方如下 ############')
        for i in self.cell_S3S4g3g4:
            print(i.gene,i.rna,i.nucleotide,i.amino_acid,i.exon,i.variant_type,i.vaf)
        print('S3S4g3g4 所有出问题的地方如上 ############')

    class Mutation_ele():
        def __init__(self,gene,rna,nucleotide,amino_acid,exon,variant_type,vaf_cnv):
            self.gene = gene
            self.rna = rna
            self.nucleotide = nucleotide
            self.amino_acid = amino_acid
            self.exon = exon
            self.variant_type = variant_type
            self.vaf = vaf_cnv
            self.replace_nan()
        def replace_nan(self):
            for attr, value in self.__dict__.items():
                if isinstance(value, float) and np.isnan(value):
                    setattr(self, attr, '')

    position=[]
    @Tools_Decorator(tool='format_Cell',table_index=0,position = position  ,size=10) # 改字体.
    @log_to_file(folder='./', filename='1.log')
    def sample_information(self, position=[]):
        for temp_table in self.doc.tables:
            if len(temp_table.columns) == 6 and temp_table.cell(0, 0).text == '基本信息':
                target_table = temp_table
                break  # 记得停下来，我们要找到额其实是第一个 这样表头的 table。
        table = target_table
        table.cell(1, 1).text = self.name
        table.cell(1, 3).text = self.gender
        table.cell(1, 5).text = self.age
        table.cell(2, 1).text = self.hospital
        table.cell(2, 3).text = self.arrival_date
        table.cell(3, 1).text = self.clinname
        table.cell(3, 3).text = self.projectname
        table.cell(4, 1).text = self.sample#.replace('-T','').replace('-N','')
        table.cell(4, 3).text = self.sample_type
        table.cell(5, 1).text = '-'
        temp_position = [(1, 1), (1, 3), (1, 5), (2, 1), (2, 3), (3, 1), (3, 3), (4, 1), (4, 3), (5, 1)]
        position.extend(temp_position)
        return self.doc

    position,position_1 = [],[]
    @Tools_Decorator(tool='format_Cell', table_index=4,position=position, size=10, color=RGBColor(255, 0, 0))
    @Tools_Decorator(tool='format_Cell', table_index=4,position=position_1, size=10, color=RGBColor(0, 0, 0))
    @log_to_file(folder='./', filename='1.log')
    def Details_of_genetic_testing_results(self, position=[],position_1=[]):  # 给第个表格插入值。
        for temp_table in self.doc.tables:
            if len(temp_table.columns) == 3 and temp_table.cell(0, 0).text == '基因' \
                and temp_table.cell(0, 1).text == '检测内容' \
                and temp_table.cell(0, 2).text == '检测结果（仅显示致病/可能致病变异）' :
                target_table = temp_table
                break
        table = target_table
        # print(len(table.rows), len(table.columns))
        for i in range(1, len(table.rows)):
            table.cell(i, 2).paragraphs[0].text = '野生型/未检出'
        # 每一个位置可能会被重复插入。\t 分割
        # 这里为了安全，本来不用exec和eval 想使用local或者global来动态生成变量。global的区域以函数分割。
        # 但是 local 动态赋值是不可取的，local 只读 global 可读可写。参考<local陷阱>：https://blog.csdn.net/weixin_43064185/article/details/104141327
        # 我又不想搞模块全局变量赋值，所以还是选用来有点危险的 exec 方法来动态赋值。小心一点。
        # 在函数a()中使用exec()函数动态创建变量时，这些变量将存在于执行exec()函数的作用域中,外部作用域无法调用，
        # 外部作用域中无法直接引用这些变量。
        # 解决这个问题的方法是将动态创建的变量存储在一个可访问的数据结构中，例如字典或列表，以便在需要时进行引用。
        # 方案一，用list，像我在制作污染库时一样。
        # self.cell = []
        # for j in range(1, len(table.rows)):
        #     # locals()[f'cell_{j}'] = '_'  # local 动态赋值不可取。
        #     exec(f"cell_{j} = '_' ")
        #     self.cell.append(locals()[f'cell_{j}'])
        # 方案二，用dict。
        cell = {}
        for j in range(1, len(table.rows)):
            gene,detection_content = table.cell(j,0).text,table.cell(j,1).text
            # locals()[f'cell_{j}'] = '_'  # local 动态赋值不可取。
            exec(f"cell['cell_{j}'] = [gene,detection_content,''] ")
        # print(cell['cell_2'])
        # 记录要保留的 index 。进入《基因检测结果详情》 表的self.cell 元素才保留。
        reserve_index_list = []
        for index,ele in enumerate(self.cell):
            # 第一先考虑snpindel 这个sheet 中的s1 情况。
            if ele.variant_type != '基因融合' and ele.variant_type != '基因扩增':
                for j in range(1, len(table.rows)):
                    if ele.gene in cell[f'cell_{j}'][0].split(' ') and ('突变' in  cell[f'cell_{j}'][1] or (f"外显子" in cell[f'cell_{j}'][1] and f"{ele.exon.replace('exon','')}" in cell[f'cell_{j}'][1])):
                        cell[f'cell_{j}'][2] =  cell[f'cell_{j}'][2] + ele.amino_acid + f' ({ele.vaf})\n'
                        reserve_index_list.append(index)
            # 第二考虑 fusion 这个sheet 中的s1 情况。
            if ele.variant_type == '基因融合':
                for j in range(1, len(table.rows)):
                    if (cell[f'cell_{j}'][0] in ele.gene or ele.gene in cell[f'cell_{j}'][0].split(' '))  and ('重排' in  cell[f'cell_{j}'][1] or ('其它' in  cell[f'cell_{j}'][1])):
                        cell[f'cell_{j}'][2] = cell[f'cell_{j}'][2] + ele.gene + f' 基因融合\n'
                        reserve_index_list.append(index)
            # 第三考虑 cnv 这个sheet 中的s1 情况。
            if ele.variant_type == '基因扩增':
                for j in range(1, len(table.rows)):
                    if (cell[f'cell_{j}'][0] in ele.gene or ele.gene in cell[f'cell_{j}'][0].split(' ')) and ('扩增' in cell[f'cell_{j}'][1] or ('其它' in  cell[f'cell_{j}'][1])):
                        cell[f'cell_{j}'][2] = cell[f'cell_{j}'][2] + f'基因扩增\n'
                        reserve_index_list.append(index)
        # 把 完工的cell 填入对应的talbe的位置。
        for i in range(1, len(table.rows)):
            if cell[f'cell_{i}'][2] != '':
                table.cell(i, 2).paragraphs[0].text = cell[f'cell_{i}'][2].strip('\n')
                position.append((i, 2))
        # 改字体。
        temp_position = [(i, j) for i in range(len(table.rows)) for j in range(len(table.columns))]
        position_1.extend(temp_position)
        # self.format_Cell(table, position=position_1, size=10, color=RGBColor(0, 0, 0))
        # self.format_Cell(table, position=position, size=10, color=RGBColor(255, 0, 0))
        # 不要忘记了 对 self.cell 的元素进行筛选，因为有的self.cell 原色可能不会出现在 《基因检测结果详情》表中，也就是不需后续处理。
        print('self.cell 中的原色最后保留的只有：',reserve_index_list)
        self.cell_raw = copy.deepcopy(self.cell)
        self.cell = [self.cell[i] for i in reserve_index_list]
        # 补丁：
        # 肺癌全阴 self.cell 中 加入EGFR/ALK野生型，
        # 肠癌全阴 self.cell 中 加入KRAS/NRAS/BRAF野生型；
        if self.cell == [] and self.cancer == '肺癌':
            ele = self.Mutation_ele(gene='EGFR/ALK', rna='', nucleotide='' \
                , amino_acid='', exon='',variant_type='野生型', vaf_cnv='')
            self.cell.append(ele)
        if self.cell == [] and self.cancer == '肠癌':
            ele = self.Mutation_ele(gene='KRAS/NRAS/BRAF', rna='', nucleotide='' \
                , amino_acid='', exon='',variant_type='野生型', vaf_cnv='')
            self.cell.append(ele)
        return self.doc

    @Tools_Decorator(tool='move_table_after', paragraph_end_with='具有明确临床意义的变异位点【Ⅰ类/Ⅱ类】')
    @Tools_Decorator(tool='format_Table', col_width=[1.8, 2.8, 2.8, 2.8, 1.64, 2.38, 1.7],Border_all=False)
    @Tools_Decorator(tool='format_Cell', position=position, size=10, color=RGBColor(0, 0, 0))
    @Tools_Decorator(tool='fill_empty_cells_with_slash')
    def Somatic_variation_results_0(self,position=[]):
        # print(self.cell)     # ['ALK EML401-ALK 基因融合', 'ALK EML400-ALK 基因融合', 'EGFR p.L747_E749del', 'EGFR p.Lzhoujielun', 'EGFR p.zhoujielun', 'EGFR p.zhoujluner', 'MAP2K1 p.liudehua', 'MET MET 基因扩增']
        # print(self.cell,'注意看这个 self cell')
        # 建立表1。
        header = ['基因', '转录本编号', '核苷酸变化', '氨基酸变化', '外显子', '变异类型', '丰度/\n拷贝数']
        if len(self.cell) == 0:  # 空 情况。
            table = self.doc.add_table(2, 7)
            for i in range(len(header)):  # 列
                table.cell(0, i).paragraphs[0].text = header[i]
            table.cell(1, 0).merge(table.cell(1, 6))
            table.cell(1, 0).paragraphs[0].text = '未检出相关基因突变'
        else:
            table = self.doc.add_table(len(self.cell) + 1, 7)
            for i in range(7):
                for j in range(len(self.cell) + 1):
                    table.cell(j, i).paragraphs[0].text = '/'
            for i in range(0, len(header)):  # 列
                table.cell(0, i).paragraphs[0].text = header[i]
            for index, ele in enumerate(self.cell):
                table.cell(index+1, 0).paragraphs[0].text = ele.gene
                table.cell(index+1, 1).paragraphs[0].text = ele.rna
                table.cell(index+1, 2).paragraphs[0].text = ele.nucleotide
                table.cell(index+1, 3).paragraphs[0].text = ele.amino_acid
                table.cell(index+1, 4).paragraphs[0].text = ele.exon.replace('exon', '')
                table.cell(index+1, 5).paragraphs[0].text = ele.variant_type
                # print('ok',ele.gene,ele.variant_type,ele.vaf)
                table.cell(index+1, 6).paragraphs[0].text = str(ele.vaf)
                # print('ok1',ele.gene,ele.variant_type,ele.vaf)
        self.table_analysis = copy.deepcopy(table)
        # self.table_analysis = copy.deepcopy(table)  # 拷贝好这个表格，给后面一个插入的地方用，记得表格的拷贝要靠 深拷贝。
        temp_position = [(i, j) for i in range(len(table.rows)) for j in range(len(table.columns))]
        position.extend(temp_position)
        return self.doc

    @Tools_Decorator(tool='move_table_after', paragraph_end_with="具有潜在临床意义的变异位点【Ⅲ类】")
    @Tools_Decorator(tool='format_Table', col_width=[1.8, 2.8, 2.8, 2.8, 1.64, 2.38, 1.7],Border_all=False)
    @Tools_Decorator(tool='format_Cell', position=position, size=10, color=RGBColor(0, 0, 0))
    def Somatic_variation_results_1(self,position=[]):
        # 建立表2。
        header = ['基因', '转录本编号', '核苷酸变化', '氨基酸变化', '外显子', '变异类型', '丰度/\n拷贝数']
        if len(self.cell_S3S4g3g4) == 0:  # 空 情况。
            table = self.doc.add_table(2, 7)
            for i in range(len(header)):
                table.cell(0, i).paragraphs[0].text = header[i]
            table.cell(1, 0).merge(table.cell(1, 6))
            table.cell(1, 0).paragraphs[0].text = '未检出相关基因突变'
        else:
            num_row = len(self.cell_S3S4g3g4)
            table = self.doc.add_table(num_row + 1, 7)
            for i in range(7):
                for j in range(num_row + 1):
                    table.cell(j, i).paragraphs[0].text = '/'
            for i in range(len(header)):
                table.cell(0, i).paragraphs[0].text = header[i]
            for index, ele in enumerate(self.cell_S3S4g3g4):
                table.cell(index+1, 0).paragraphs[0].text = ele.gene
                table.cell(index+1, 1).paragraphs[0].text = ele.rna
                table.cell(index+1, 2).paragraphs[0].text = ele.nucleotide
                table.cell(index+1, 3).paragraphs[0].text = ele.amino_acid
                table.cell(index+1, 4).paragraphs[0].text = ele.exon.replace('exon', '')
                table.cell(index+1, 5).paragraphs[0].text = ele.variant_type
                table.cell(index+1, 6).paragraphs[0].text = ele.vaf
        temp_position = [(i, j) for i in range(len(table.rows)) for j in range(len(table.columns))]
        position.extend(temp_position)
        # self.cell_0 = copy.deepcopy(list(map(lambda x: x.strip(' '), self.cell)))
        # # print(self.cell_1,self.cell_0)
        # # 插入表。
        # target_paragraph = self.fixed_position('Somatic_variation_results_table')
        # # print(target_paragraph.text)
        # table_before_paragraph_0 = target_paragraph.insert_paragraph_before('◆ 致病或可能致病的变异位点【Ⅰ类/Ⅱ类】')
        # self.format_paragraph(table_before_paragraph_0, size=11, bold=True)
        # table_before_paragraph_1 = target_paragraph.insert_paragraph_before("\n◆ 临床意义不明确的变异位点【Ⅲ类】")
        # self.format_paragraph(table_before_paragraph_1, size=11, bold=True)
        # self.move_table_after(table, table_before_paragraph_0)
        # self.move_table_after(table_1, table_before_paragraph_1)
        # # 改字体。
        # position = [(i, j) for i in range(len(table.rows)) for j in range(len(table.columns))]
        # self.format_Cell(table, position=position, size=10, color=RGBColor(0, 0, 0))
        # position = [(i, j) for i in range(len(table_1.rows)) for j in range(len(table_1.columns))]
        # self.format_Cell(table_1, position=position, size=10, color=RGBColor(0, 0, 0))
        # # 改表格。
        # self.format_Table(table, col_width=[2, 3, 2.7, 2.8, 1.5, 2.3, 1.7], Border=[len(table.rows) - 1])
        # self.format_Table(table_1, col_width=[2, 3, 2.7, 2.8, 1.5, 2.3, 1.7], Border=[len(table_1.rows) - 1])
        return self.doc

    def targeted_Therapy_Tips(self,position=[]):
        # 构建靶向治疗提示表 的规则方法。
        def build_Rules(Genetic_variation, variation_type) -> (str, str, str, str):  # 一次处理一行。返回的是  # 本癌药物 、其它癌药物、其它癌药物、耐药药物。
            ele = Genetic_variation.split('\n')
            df = pd.read_excel('/refhub/ref/var/VarLabel_Drug_Sensitivity.xlsx', sheet_name=None)
            df_sensitive = df['实体瘤用药表']
            # 按照snpdelins 方法来
            if re.findall(r'基因融合', Genetic_variation) == [] and re.findall(r'基因扩增', Genetic_variation) == [] and re.findall(r'野生型', Genetic_variation) == []:
                name = ele[0]
                alteration_list = []
                alteration_0 = ele[1].replace(r'p.', '')  # L747_E749del
                alteration_1_list = re.findall(r'\d+', alteration_0)  # ['747', '749']
                alteration_2 = '致病突变'
                alteration_3 = variation_type
                alteration_list.extend([alteration_0, alteration_2, alteration_3])
                # print(name,alteration_0,alteration_1_list,alteration_2,alteration_3)
                df_result_0 = df_sensitive[
                    (df_sensitive['Gene'] == name) & (df_sensitive['VarLabel'].isin(alteration_list))]
                # df['col1'].astype(str).str.findall('\d+').apply(set).apply(lambda x: bool(x & set(a)))
                # x & set(a)是一个交集运算，它返回两个集合的交集。如果交集不为空，则返回True，否则返回False。
                # 因此，apply(lambda x: bool(x & set(a)))的作用是判断数字集合是否与列表a有交集，如果有，则返回True，否则返回False。
                df_result_1 = df_sensitive[(df_sensitive['Gene'] == name) & (
                    df_sensitive['VarLabel'].astype(str).str.findall('\d+').apply(set).apply(
                        lambda x: bool(x & set(alteration_1_list))))]
                df_result = pd.concat([df_result_0, df_result_1])
                df_result.drop_duplicates(keep='first', inplace=True)  # 所有列都相等就去重。# df_result_1 df_result_0 之间可能有重复的部分。
                # print(df_result)
                # 按照fusion 方法来
            elif re.findall(r'基因融合', Genetic_variation) != []:
                name_list = ele[0].split('-')
                alteration_0 = '融合'
                df_result_0 = df_sensitive[
                    (df_sensitive['Gene'] == name_list[0]) & (df_sensitive['VarLabel'].isin([alteration_0]))]
                df_result_1 = df_sensitive[
                    (df_sensitive['Gene'] == name_list[1]) & (df_sensitive['VarLabel'].isin([alteration_0]))]
                df_result = pd.concat([df_result_0, df_result_1])
                # print(name_list, alteration_0)
                # print(df_result)
            # 按照cnv 方法来
            elif re.findall(r'基因扩增', Genetic_variation) != []:
                name = ele[0]
                alteration_0 = '扩增'
                df_result = df_sensitive[
                    (df_sensitive['Gene'] == name) & (df_sensitive['VarLabel'].isin([alteration_0]))]
                # print(name,alteration_0)
                # print(df_result)
            elif re.findall(r'野生型', Genetic_variation) != []:
                name = ele[0]
                alteration_0 = '野生型'
                df_result = df_sensitive[
                    (df_sensitive['Gene'] == name) & (df_sensitive['VarLabel'].isin([alteration_0]))]
            # df_result = df_result[df_result['Sensitivity'] == '敏感']
            # print(df_result)
            drug_list_raw = df_result['Drug'].tolist()
            drug_list = []
            for items in drug_list_raw: drug_list.extend(
                items.split(','))  # ['奥希替尼', '伏美替尼', '厄洛替尼+贝伐珠单抗', '纳武单抗+厄洛替尼']
            drug_list = list(set(drug_list))
            # print(drug_list)
            self_drug, other_drug, Clinical_drugs, resistance_Drug = '/', '/', '/', '/'
            for i, row in df_result.iterrows():
                # print(row)
                for drug in row['Drug'].split(','):
                    # print(drug)
                    # 临床药物
                    if self.DrugApproval[self.DrugApproval['DrugName'].isin([drug])].empty:
                        if row['Sensitivity'] == '敏感': Clinical_drugs += drug + ','
                        # print('临床药物: ',Clinical_drugs)
                    # 本癌药物,潜在耐药药物。(定义耐药药物是指，由这个变异搜到的药物匹配癌种对应相同，但是变异对药物有耐药性。)
                    elif not self.DrugApproval[self.DrugApproval['DrugName'].isin([drug])].empty:
                        # 加入一个 审批补丁方法，有的药物只针对固定的突变 通过了审批，
                        # 如果没有通过此补丁，相当于此药物未通过审批,回到上面看看能不能算成临床药物。
                        def review_Patches(drug, alteration_0) -> bool:  # 这里我写的有点不一样，方便后续拓展，只是读起来有点绕。
                            options = {('索托拉西布'): 'G12C', ('阿达格拉西布'): 'G12C'}
                            return options.get((drug), alteration_0) != alteration_0

                        # print('look here!!! ',drug,alteration_0)
                        if review_Patches(drug, alteration_0):  # 打补丁
                            # 临床药物
                            if row['Sensitivity'] == '敏感': Clinical_drugs += drug + ','
                        else:  # 不打补丁
                            df_temp = self.DrugApproval[
                                self.DrugApproval['DrugName'].isin([drug])]  # 这个 df_temp 是 drug approve表中的 子集。
                            # print(df_temp)
                            # 和 上面一样，使用求交集的运算，更加灵活一些。下面这个print 调试时候用。
                            # print(f'df_temp {drug} 的匹配： ',df_temp[df_temp['OncoType'].astype(str).str.split(',').apply(set).apply(
                            #     lambda x: bool(x & set(self.BC17_meta['cancer'].tolist())))])
                            if not df_temp[df_temp['OncoType'].astype(str).str.split('[,;；，]').apply(set).apply(
                                    lambda x: bool(x & set(self.BC17_meta['cancer'].tolist())))].empty:
                                if row['Sensitivity'] == '敏感':  self_drug += drug + ','
                                if row['Sensitivity'] == '耐药': resistance_Drug += drug + ','
                            # print('本癌药物: ', self_drug)
                            # 其它癌药物
                            if df_temp[df_temp['OncoType'].astype(str).str.split(',').apply(set).apply(
                                    lambda x: bool(x & set(self.BC17_meta['cancer'].tolist())))].empty:
                                if row['Sensitivity'] == '敏感':  other_drug += drug + ','
                                # print('其它癌药物: ', other_drug)
            # 去重，处理格式。
            if len(self_drug) > 1: self_drug = '\n'.join(list(set(self_drug.strip('/').strip(',').split(','))))
            if len(other_drug) > 1: other_drug = '\n'.join(list(set(other_drug.strip('/').strip(',').split(','))))
            if len(Clinical_drugs) > 1: Clinical_drugs = '\n'.join(
                list(set(Clinical_drugs.strip('/').strip(',').split(','))))
            if len(resistance_Drug) > 1: resistance_Drug = '\n'.join(
                list(set(resistance_Drug.strip('/').strip(',').split(','))))
            # 注意:
            # 以下为 是否生成 中间文件，可选。(第一次运行的时候要记得打开，生成的中间文件后续会用上。)
            if not os.path.isfile(f'./{self.sample}药物注释匹配表.csv'):
                df_result.to_csv(f'./{self.sample}药物注释匹配表.csv', mode='a', index=False)
            else:
                df_result.to_csv(f'./{self.sample}药物注释匹配表.csv', mode='a', header=False, index=False)
            if not os.path.isfile(f'./{self.sample}注释匹配表.csv'):
                citation_list = []
                Citations = df_result['Citation'].tolist()
                for items in Citations: citation_list.extend(str(items).split(','))
                with open(f'./{self.sample}注释匹配表.csv', 'w')as f:
                    f.write(Genetic_variation.replace('\n', ' ') + ':' + ','.join(citation_list) + '\n')
            else:
                citation_list = []
                Citations = df_result['Citation'].tolist()
                for items in Citations: citation_list.extend(str(items).split(','))
                with open(f'./{self.sample}注释匹配表.csv', 'a+')as f:
                    f.write(Genetic_variation.replace('\n', ' ') + ':' + ','.join(citation_list) + '\n')
            # print('self_drug:', self_drug, 'other_drug:', other_drug, 'Clinical_drugs:', Clinical_drugs,'resistance_Drug: ', resistance_Drug)
            return self_drug, other_drug, Clinical_drugs, resistance_Drug

        @Tools_Decorator(tool='move_table_after', paragraph_end_with="靶向治疗提示")
        @Tools_Decorator(tool='format_Table', col_width=[3, 3, 3.92, 3, 3])
        @Tools_Decorator(tool='format_Cell', position=position, size=10, color=RGBColor(0, 0, 0))
        def have_targeted_Therapy_Tips(position=[]):
            # 建立表。
            table = self.doc.add_table(len(self.cell) + 1, 5)
            for i in range(5):
                for j in range(len(self.cell) + 1):
                    table.cell(j, i).paragraphs[0].text = '/'
            header = ['基因变异', '推荐本癌种药物', '推荐其他癌种药物', '临床试验药物', '潜在耐药药物']
            for i in range(len(header)):
                table.cell(0, i).paragraphs[0].text = header[i]
            for index,ele in enumerate(self.cell):
                # print(row.cells[5].text,re.findall(r'基因融合',row.cells[5].text)==[])
                # if i > 0:
                #     if re.findall(r'基因融合', row.cells[5].text) == [] and re.findall(r'基因扩增', row.cells[5].text) == []:
                #         Genetic_variation = str(row.cells[0].text) + '\n' + str(row.cells[3].text)
                #     elif re.findall(r'基因融合', row.cells[5].text) != [] or re.findall(r'基因扩增', row.cells[5].text) != []:
                #         Genetic_variation = str(row.cells[0].text) + '\n' + str(row.cells[5].text)
                #     table.cell(i, 0).text = Genetic_variation  # "EML4-ALK 基因融合" or "EGFR p.L747_E749del" or "MET 基因扩增"
                if ele.variant_type=='基因融合' or ele.variant_type=='基因扩增':
                    Genetic_variation = ele.gene + '\n' + ele.variant_type
                else:
                    Genetic_variation = ele.gene + '\n' + ele.amino_acid
                self_drug, other_drug, Clinical_drugs, resistance_Drug = build_Rules(Genetic_variation,ele.variant_type)
                table.cell(index+1, 0).text = ele.gene+'\n'+ele.amino_acid if ele.variant_type not in ['基因融合','基因扩增','野生型'] else ele.gene+'\n'+ele.variant_type
                table.cell(index+1, 1).text = self_drug
                table.cell(index+1, 2).text = other_drug
                table.cell(index+1, 3).text = Clinical_drugs
                table.cell(index+1, 4).text = resistance_Drug
                # temp_df = self.DrugApproval.loc[self.DrugApproval['ApprovedContent'].apply(lambda x: re.compile(genename).search(x) is not None)]
                # temp = '\n'.join(temp_df.drop_duplicates(subset=['DrugName'], keep='first')['DrugName'].tolist())
                # table.cell(j, 1).paragraphs[0].text = temp if temp != '' else '/'
            self.Drug_Annotation_Matching_Table_df = pd.read_csv(f'./{self.sample}药物注释匹配表.csv',sep=',')
            self.Annotation_Matching_Table_df = pd.read_csv(f'./{self.sample}注释匹配表.csv', sep=':',header=None,names=['Genetic_variation','Citations'])
            print(f'{self.sample}药物注释匹配表 is :',self.Drug_Annotation_Matching_Table_df)
            print(f'{self.sample}注释匹配表 is :', self.Annotation_Matching_Table_df)
            for file_path in [f'./{self.sample}药物注释匹配表.csv',f'./{self.sample}注释匹配表.csv']:
                if os.path.exists(file_path):
                    os.remove(file_path)
            temp_position = [(i, j) for i in range(len(table.rows)) for j in range(len(table.columns))]
            position.extend(temp_position)
            # 提取表格数据
            data = []
            for row in table.rows:
                data.append([cell.text for cell in row.cells])
            # 创建DataFrame
            self.targeted_Therapy_Tips_df = pd.DataFrame(data[1:], columns=data[0])
            return self.doc

        @Tools_Decorator(tool='move_table_after', paragraph_end_with="靶向治疗提示")
        @Tools_Decorator(tool='format_Table', col_width=[3.92, 3, 3, 3, 3],Border=[1])
        @Tools_Decorator(tool='format_Cell', position=position, size=10, color=RGBColor(0, 0, 0))
        def not_have_targeted_Therapy_Tips(position=[]):
            # 建立表。
            table = self.doc.add_table(2, 5)
            header = ['基因变异', '推荐本癌种药物', '推荐其他癌种药物', '临床试验药物', '潜在耐药药物']
            for i in range(len(header)):  # 列
                table.cell(0, i).paragraphs[0].text = header[i]
            table.cell(1, 0).merge(table.cell(1, 4))
            table.cell(1, 0).paragraphs[0].text = '未检出相关基因突变，无相关提示。'
            temp_position = [(i, j) for i in range(len(table.rows)) for j in range(len(table.columns))]
            position.extend(temp_position)
            data = {
                '基因': ['/'],
                '推荐本癌种药物': ['/'],
                '推荐其他癌种药物': ['/'],
                '临床试验药物': ['/'],
                '潜在耐药药物': ['/']
            }
            self.targeted_Therapy_Tips_df = pd.DataFrame(data)
            return self.doc

        if len(self.cell) != 0:
            doc = have_targeted_Therapy_Tips(position=[])
            return doc
        else:
            doc = not_have_targeted_Therapy_Tips(position=[])
            return doc

    def Analysis_of_Somatic_Variant_Genes_and_Loci(self,position=[]):  # 模版中 《体细胞变异基因及位点解析》《靶向药物注释》之间是没有分页符的。
        def Add_clinical_significance_tips_to_sub_table(row_index) -> str:  # 此方法会在 add_sub_table 方法内部被调用。
            # df_temp = pd.read_csv(f'/refhub/ref/matching/{self.sample}注释匹配表.csv', header=None,
            #                       names=['Genetic_variation', 'Citations'],
            #                       sep=':')  # 读前面add_table_Targeted_treatment_tips生成的中间表。
            df_temp = self.Annotation_Matching_Table_df
            # 如果 citations 没有匹配到。
            df_temp['Citations'].fillna('', inplace=True)
            # print(df_temp)
            # print('df_temp["Citations"].iloc[row_index]',df_temp['Citations'].iloc[row_index])
            # print("df_temp['Citations'].iloc[row_index],row_index,df_temp['Citations']",df_temp['Citations'].iloc[row_index],row_index,df_temp['Citations'])
            citation_list = df_temp['Citations'].astype(str).iloc[row_index].split(',')
            # citation_list = df_temp['Citations'].iloc[row_index].split(',')
            while 'nan' in citation_list:  # 去掉 citation 中 nan 的值。
                citation_list.remove('nan')
            description_list = []
            citation_list = list(set(citation_list))
            # 把 特定的 元素 提前。这里的规则是
            # 左侧优先排CSCO中国标准，其次是NCCN的美国标准，最后是数字类型的。
            citation_list = [i for i in citation_list if re.findall('csco', i) != []] \
                            + [i for i in citation_list if re.findall('nccn', i) != []] \
                            + [i for i in citation_list if re.findall('[a-z]+', i) == []]
            # citation_list = [i for i in citation_list if re.findall('[a-z]+', i) != []] + [i for i in citation_list if re.findall('[a-z]+',i) == []]
            # print(citation_list)
            if citation_list == ['']:
                return '无'
            for citation in citation_list:
                # print('citation is :',citation)
                # 注意，可能出现 citation 从 citation表中找不到的情况，此时 会找到 空的 Series([], Name: Description, dtype: object)
                if self.Citation[self.Citation['Citation'].astype(str) == citation.strip(' ')]['Description'].empty:
                    print('出现 citation 从 citation表中找不到的情况,跳过。')
                    continue

                # print(self.Citation[self.Citation['Citation'].astype(str) == citation.strip(' ')]['Description'])
                # 加入一个 citation补丁方法，针对特定的基因，如果不是特定的突变，且description中有特定突变字眼,就不把这段description 加进去。
                def citation_Patches(gene, mutation):  # 这里我写的有点不一样，方便后续拓展，只是读起来有点绕。
                    options = {('KRAS'): 'G12C'}
                    return options.get((gene), mutation) != mutation, options.get((gene), mutation)

                [gene, mutation] = df_temp['Genetic_variation'].iloc[row_index].split(' ')
                mutation = mutation.replace('p.', '')
                if citation_Patches(gene, mutation)[0]:  # 打补丁,就是跳过特定description 的添加。
                    description = str(
                        self.Citation[self.Citation['Citation'].astype(str) == citation.strip(' ')]['Description'].iloc[
                            0])
                    options_value = str(citation_Patches(gene, mutation)[1])
                    if re.findall(options_value, description) != []:
                        continue
                    description_list.append(
                        self.Citation[self.Citation['Citation'].astype(str) == citation.strip(' ')]['Description'].iloc[
                            0])
                else:  # 不打补丁
                    description_list.append(
                        self.Citation[self.Citation['Citation'].astype(str) == citation.strip(' ')]['Description'].iloc[
                            0])
            # print(description_list,'look here!')
            description_list = list(filter(lambda x: isinstance(x, str), description_list))
            return '\n'.join(description_list)

        position,position_1,position_0 = [],[],[]
        @Tools_Decorator(tool='format_Table', col_width=[1.8, 2.8, 2.8, 2.8, 1.64, 2.38, 1.7], Border_all=True)
        @Tools_Decorator(tool='format_Cell', position=position_1, size=10, color=RGBColor(0, 0, 0), left=False)
        @Tools_Decorator(tool='format_Cell', position=position, size=10, color=RGBColor(0, 0, 0),left=True)
        def add_sub_table(i,position=[],position_1=[],position_0=[]):  # address[1] 代表 了 体细胞变异结果 表的第几行,1只是默认值，可以变的。# 每次新建一个table 放最后。
            addresss = [0, i]
            # 在体细胞变异结果表的基础上，选取几行构建新的表，再在此基础上添加行，构建新的表。
            new_table = self.doc.add_table(rows=1, cols=len(self.table_analysis.columns))  # 先创建一个一行的表。
            for i, row in enumerate(self.table_analysis.rows):
                if i in addresss:  # 这里取第0 和 第某 行 构建新的表。
                    new_row = new_table.add_row()  # 创建一个新的行
                    for j, cell in enumerate(row.cells):
                        new_row.cells[j].text = cell.text if cell.text else '/' # 110 80 也要改
            rows = new_table.rows  # 删除第一行的空白行。
            new_table._element.remove(rows[0]._element)
            for i in range(6):  # 每个新的表下面再加上6行。
                new_table.add_row()
                new_table.cell(i + 2, 0).merge(new_table.cell(i + 2, len(self.table_analysis.columns) - 1))
                if i == 0:
                    new_table.cell(i + 2, 0).paragraphs[0].text = '基因与肿瘤相关性概述'
                if i == 1:
                    describe = self.GeneIndo[
                        self.GeneIndo['Genename'] == self.table_analysis.cell(addresss[1], 0).paragraphs[0].text][
                        'Describe']
                    # print(describe)
                    new_table.cell(i + 2, 0).paragraphs[0].text = describe.astype(str).iloc[0] if (
                                                                                                      not describe.empty) and (
                                                                                                      not describe.isna().any()) else '/'
                if i == 2:
                    new_table.cell(i + 2, 0).paragraphs[0].text = '位点变异信息注释'
                if i == 3:
                    fill_in_00 = self.table_analysis.cell(addresss[1], 0).paragraphs[0].text
                    fill_in_0 = self.table_analysis.cell(addresss[1], 0).paragraphs[0].text + " " + \
                                self.table_analysis.cell(addresss[1], 3).paragraphs[0].text
                    fill_in_1 = self.table_analysis.cell(addresss[1], 5).paragraphs[0].text
                    fill_in_2 = self.table_analysis.cell(addresss[1], 0).paragraphs[0].text
                    fill_in_3 = self.table_analysis.cell(addresss[1], 4).paragraphs[0].text
                    fill_in_4 = self.table_analysis.cell(addresss[1], 3).paragraphs[0].text
                    fill_in_5 = self.table_analysis.cell(addresss[1], 6).paragraphs[0].text
                    # for row in self.table_analysis.rows:
                    #     for cell in row.cells:
                    #         print(cell.text,end='')
                    #     print()
                    # print('fill_in_4 is : ',fill_in_4)
                    # print('fill_in_1 is : ',fill_in_1)
                    if fill_in_4 != '/':
                        e = explain()
                        e.Interpretative(fill_in_4)
                        temp_text = f'患者样本中检出的{fill_in_0}为{fill_in_1},该突变位于{fill_in_2}基因的第{fill_in_3}号外显子,' \
                                    f'{e.Interpretative(fill_in_4)}'
                        # temp_text = 'ok'
                    else:
                        if fill_in_1 == '基因扩增':
                            temp_text = f"患者样本中检出的{fill_in_00}为{fill_in_1}变异,该突变检测到{fill_in_5}个拷贝数扩增,导致蛋白质表达水平的变化。"
                        elif fill_in_1 == '基因融合':
                            temp_text = f"患者样本中检出的{fill_in_00}为{fill_in_1}变异,该突变检测到{fill_in_00.split('-')[0]}和{fill_in_00.split('-')[1]}发生了融合,导致蛋白质表达水平的变化。"
                    new_table.cell(i + 2, 0).paragraphs[0].text = temp_text
                if i == 4:
                    new_table.cell(i + 2, 0).paragraphs[0].text = '临床意义提示'
                if i == 5:
                    fill_in_6 = Add_clinical_significance_tips_to_sub_table(addresss[1] - 1)
                    new_table.cell(i + 2, 0).paragraphs[0].text = fill_in_6
            temp_position = [(i, j) for i in range(len(new_table.rows)) for j in range(len(new_table.columns))]
            position.extend(temp_position)
            position_1.extend([(0,0),(0,1),(0,2),(0,3),(0,4),(0,5),(0,6),(1,0),(1,1),(1,2),(1,3),(1,4),(1,5),(1,6)])
            position_0.extend([(3,0),(5,0),(7,0)])
            # 微改表格。
            for i in [3, 5, 7]:
                for paragraph in new_table.cell(i, 0).paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            new_table.rows[7].height = Cm(5)
            return self.doc
        def insert_add_sub_table(i): # new_table 直接用最后一个 table。
            address = [0, i]
            add_sub_table(i,position=[],position_1=[],position_0=[])
            # new_table = self.doc.add_table(2, 5)
            # 插入表。(但是注意，由于不能在段落后插入分页符，只能先定位到后一个段落，再在后一个段落前插入分页符。下面这样写是为了格式布局。模版中的体细胞变异基因及位点解析 后的分页符先删除。)
            target_paragraph = self.class_Tools.fixed_position(self.doc,End_character='体细胞变异基因及位点解析')
            # print('address[1]',address[1],'len(self.cell)',len(self.cell))
            if address[1] == 1 and address[1] != len(self.cell):  # 第一个子表 且不只有一个子表。
                # print('走的是第一条路。')
                self.class_Tools.move_table_after(self.doc, paragraph_end_with='体细胞变异基因及位点解析') # doc, paragraph_end_with,table_index=-1
            elif address[1] == 1 and address[1] == len(self.cell):  # 只有一个子表。
                # print('走的是第二条路。')
                self.class_Tools.move_table_after(self.doc, paragraph_end_with='体细胞变异基因及位点解析')
                next_paragraph = self.class_Tools.fixed_position(self.doc,End_character='靶向药物注释')
                paging_paragraph = next_paragraph.insert_paragraph_before(
                    '')  # (f"{address[1]}'")  # 分页符段落没法定位，所以只能在插入的时候赶紧完成后续操作。
                run = paging_paragraph.add_run()
                run.add_break(WD_BREAK.PAGE)
            elif address[1] == len(self.cell) and address[1] != 1:  # 最后一个子表。
                # print('走的是第三条路。')
                next_paragraph = self.class_Tools.fixed_position(self.doc,End_character='靶向药物注释')
                # print(next_paragraph)
                paging_paragraph = next_paragraph.insert_paragraph_before(
                    '')  # (f"{address[1]}'")  # 分页符段落没法定位，所以只能在插入的时候赶紧完成后续操作。
                run = paging_paragraph.add_run()
                run.add_break(WD_BREAK.PAGE)
                self.class_Tools.move_table_after_supplement(paging_paragraph)
                paging_paragraph1 = next_paragraph.insert_paragraph_before('')  # (f"{address[1]}'")
                run = paging_paragraph1.add_run()
                run.add_break(WD_BREAK.PAGE)
            else:
                # print('走的是第四条路。')
                next_paragraph = self.class_Tools.fixed_position(self.doc,End_character='靶向药物注释')
                # print(next_paragraph)
                paging_paragraph = next_paragraph.insert_paragraph_before(
                    '')  # (f"{address[1]}'")  # 分页符段落没法定位，所以只能在插入的时候赶紧完成后续操作。
                run = paging_paragraph.add_run()
                run.add_break(WD_BREAK.PAGE)
                self.class_Tools.move_table_after_supplement(paging_paragraph)


        @Tools_Decorator(tool='move_table_after', paragraph_end_with='体细胞变异基因及位点解析')
        @Tools_Decorator(tool='format_Table',col_width=[2, 3, 2.7, 2.8, 1.5, 2.3, 1.7],  Border=[1])
        @Tools_Decorator(tool='format_Cell', position=position, size=10, color=RGBColor(0, 0, 0))
        def not_insert_add_sub_table(position=[]):
            # 建立表。
            new_table = self.doc.add_table(rows=1, cols=len(self.table_analysis.columns))  # 先创建一个一行的表。
            for i, row in enumerate(self.table_analysis.rows):
                if i == 0:  # 这里取第0 和 第1 行 构建新的表。
                    new_row = new_table.add_row()  # 创建一个新的行
                    for j, cell in enumerate(row.cells):
                        new_row.cells[j].text = cell.text
            rows = new_table.rows  # 删除第一行的空白行。
            new_table._element.remove(rows[0]._element)
            new_table.add_row()
            new_table.cell(1, 0).merge(new_table.cell(1, 6))
            new_table.cell(1, 0).text = '未检出相关基因突变，无相关提示。'
            new_table.cell(1, 0).merge(new_table.cell(1, 6))
            next_paragraph = self.class_Tools.fixed_position(self.doc,End_character='靶向药物注释')
            paging_paragraph1 = next_paragraph.insert_paragraph_before('')  # (f"{address[1]}'")
            run = paging_paragraph1.add_run()
            run.add_break(WD_BREAK.PAGE)
            temp_position = [(i, j) for i in range(len(new_table.rows)) for j in range(len(new_table.columns))]
            position.extend(temp_position)
            # print(position)
            return self.doc

        if self.cell != []:
            for i in range(1, len(self.cell) + 1):
                doc = insert_add_sub_table(i)
        elif self.cell == []:
            doc = not_insert_add_sub_table(position=[])
        return doc

    @Tools_Decorator(tool='format_Cell', table_index=-7,position=position, size=10, color=RGBColor(0, 0, 0))
    def quality_Control_Results(self,position=[]):  # 插入值完成填表。。
        # print('进入质控表。')
        # 通过匹配table的表头来定位table。
        for index,table in enumerate(self.doc.tables):
            if len(table.columns) >= 4 and table.cell(0, 0).paragraphs[0].text == '质控内容' and \
                    table.cell(0, 1).paragraphs[0].text == '质控参数':
                target_table = table
                temp_table_index = index
        # 修改表的内容。
        # print('HE结果' in self.BC17_meta.columns)
        change_list_index = [1, 2, 4, 5, 6, 7, 8, 9, 10]
        self.temp = {}
        for i in change_list_index:
            self.temp[f'cell_{i}'] = ''
        if 'HE结果' in self.BC17_meta.columns:  # 假定 self.BC17_meta['HE结果'].iloc[0] 是个百分数。
            if float(self.BC17_meta['HE结果'].iloc[0].strip('%')) / 100.0 >= 0.1:
                self.temp['cell_1'] += f"{'{:.0%}'.format(float(self.BC17_meta['HE结果'].iloc[0].strip('%')) / 100)}"
            else:
                self.temp['cell_1'] += '>10%'
        else:
            if self.BC17_meta['样本类型*'].iloc[0] == '组织':
                self.temp['cell_1'] += '>10%'
            else:
                self.temp['cell_1'] += '不适用'
        if 'DNA总量' in self.BC17_meta.columns:
            if self.BC17_meta['样本类型*'].iloc[0] == '组织':
                if self.BC17_meta['DNA总量'].iloc[0] >= 100:
                    self.temp['cell_2'] += f"{str(int(self.BC17_meta['DNA总量'].iloc[0]))}"
                else:
                    self.temp['cell_2'] += f"100"
            else:
                if self.BC17_meta['DNA总量'].iloc[0] >= 30:
                    self.temp['cell_2'] += f"{str(int(self.BC17_meta['DNA总量'].iloc[0]))}"
                else:
                    self.temp['cell_2'] += f"30"
        else:
            if self.BC17_meta['样本类型*'].iloc[0] == '组织':
                self.temp['cell_2'] += '100'
            else:
                self.temp['cell_2'] += '30'
        if '预文库总量' in self.BC17_meta.columns:
            if self.BC17_meta['DNA总量'].iloc[0] >= 500:
                self.temp['cell_4'] += f"{str(int(self.BC17_meta['预文库总量'].iloc[0]))}"
            else:
                self.temp['cell_4'] += '500'
        else:
            self.temp['cell_4'] += '500'
        if 'depth' in self.BC17_qc.columns:
            if self.BC17_meta['样本类型*'].iloc[0] == '组织':
                if self.BC17_qc['depth'].iloc[0] >= 1000:
                    self.temp['cell_5'] += f"{str(int(self.BC17_qc['depth'].iloc[0]))}"
                else:
                    self.temp['cell_5'] += '1000'
            else:
                if self.BC17_qc['depth'].iloc[0] >= 5000:
                    self.temp['cell_5'] += f"{str(int(self.BC17_qc['depth'].iloc[0]))}"
                else:
                    self.temp['cell_5'] += '5000'
        else:
            if self.BC17_meta['样本类型*'].iloc[0] == '组织':
                self.temp['cell_5'] += '1000'
            else:
                self.temp['cell_5'] += '5000'
        if '文库多样性' in self.BC17_meta.columns:
            if float(self.BC17_meta['文库多样性'].iloc[0].strip('%')) / 100.0 >= 0.1:
                self.temp['cell_6'] += f"{'{:.2%}'.format(float(self.BC17_meta['文库多样性'].iloc[0].strip('%')) / 100)}"
            else:
                self.temp['cell_6'] += '10%'
        else:
            self.temp['cell_6'] += '10%'
        if self.BC17_qc['inser_size_peak'].iloc[0] <= 300:
            self.temp['cell_7'] += f"{str(int(self.BC17_qc['inser_size_peak'].iloc[0]))}"
        else:
            self.temp['cell_7'] += '200'
        if float(self.BC17_qc['coverage'].iloc[0].strip('%')) / 100.0 >= 0.9:
            self.temp['cell_8'] += f"{'{:.2%}'.format(float(self.BC17_qc['coverage'].iloc[0].strip('%')) / 100)}"
        else:
            self.temp['cell_8'] += '90.05%'
        if float(self.BC17_qc['mapped'].iloc[0].strip('%')) / 100.0 >= 0.95:
            self.temp['cell_9'] += f"{'{:.2%}'.format(float(self.BC17_qc['mapped'].iloc[0].strip('%')) / 100)}"
        else:
            self.temp['cell_9'] += '95.05%'
        if float(self.BC17_qc['base_quality'].iloc[0].strip('%')) / 100.0 >= 0.8:
            self.temp['cell_10'] += f"{'{:.2%}'.format(float(self.BC17_qc['base_quality'].iloc[0].strip('%')) / 100)}"
        else:
            self.temp['cell_10'] += '80.05%'
        target_table.cell(1, 2).text = self.temp['cell_1']
        target_table.cell(2, 2).text = self.temp['cell_2']
        target_table.cell(4, 2).text = self.temp['cell_4']
        target_table.cell(5, 2).text = self.temp['cell_5']
        target_table.cell(6, 2).text = self.temp['cell_6']
        target_table.cell(7, 2).text = self.temp['cell_7']
        target_table.cell(8, 2).text = self.temp['cell_8']
        target_table.cell(9, 2).text = self.temp['cell_9']
        target_table.cell(10, 2).text = self.temp['cell_10']
        temp_position = [(i, j) for i in range(1, 12) for j in range(1,4)]
        position.extend(temp_position)
        return self.doc

    def machining_table_1(self, num_th=1):
        table = self.doc.tables[num_th]
        if len(self.cell)!=1 :
            variance_num = len(self.cell)
        elif self.cell[0].variant_type!='野生型':
            variance_num = len(self.cell)
        else:
            variance_num = 0
        text = f'检出 {variance_num} 个 Ⅰ 类致病/ Ⅱ 类可能致病突变， {len(self.cell_S3S4g3g4)} 个Ⅲ类临床意义不明突变。'
        cell = table.cell(2, 1)
        cell.text = ''
        text = text.split(' ')  # 给指定的 部分一个单独的run 结构，这个单独的run 结构我们给它标红。
        for i in range(len(text)):
            run = cell.paragraphs[0].add_run(text[i])
            if i == 1:
                run.font.color.rgb = RGBColor(255, 0, 0)
        cell = table.cell(3, 1)
        cell.text = ''
        # print(self.cell)
        main_reslut = [ele.gene + ' ' + (ele.amino_acid if ele.variant_type not in ['野生型','基因融合','基因扩增'] else ele.variant_type) for ele in self.cell]
        # print('main_reslut',main_reslut)
        for i, ele in enumerate(main_reslut):
            main_reslut[i] = chr(9312 + i) + ' ' + main_reslut[i]
        # print(main_reslut)
        result = ''.join([item + '\n' if (i + 1) % 3 == 0 else item + '  ' for i, item in enumerate(main_reslut)])
        cell.text = result.strip('\n') if (result := result.strip('\n')) else '未检出靶药相关基因突变，考虑免疫药等其他治疗方案'# '  '.join(main_reslut) if main_reslut!=[] else '无' # 110 80 也要改。
        # cell.text = ' ' + cell.text
        # 通过匹配table的表头来定位table。
        for temp_table in self.doc.tables:
            if len(temp_table.columns) == 5 and temp_table.cell(0, 0).paragraphs[0].text == '基因变异' and \
                    temp_table.cell(0, 1).paragraphs[0].text == '推荐本癌种药物':
                target_table = temp_table
        A_drug, C_drug, D_drug, Drug_resistance = [], [], [], []
        for i in range(1, len(target_table.rows)):
            A_drug.extend(target_table.cell(i, 1).text.split('\n'))
            C_drug.extend(target_table.cell(i, 2).text.split('\n'))
            D_drug.extend(target_table.cell(i, 3).text.split('\n'))
            Drug_resistance.extend(target_table.cell(i, 4).text.split('\n'))
        eles = [A_drug, C_drug, D_drug, Drug_resistance]
        for i, ele in enumerate(eles):  # 注意，在Python中，filter()函数返回的是一个迭代器，而不是一个列表。当你对一个可变对象进行操作时，它会改变原对象的值。但是，当你对一个可变对象进行过滤时，它会返回一个新的迭代器，而不是改变原对象的值。
            ele = list(filter(lambda x: x != '/', list(set(ele))))
            ele = list(map(lambda x:'-' if '未检出' in x else x,ele)) # 110 80 也要改
            if i == 0: A_drug = list(map(lambda x: x + '【A级】' if '-' not in x else '', ele))
            if i == 1: C_drug = list(map(lambda x: x + '【C级】' if '-' not in x else '', ele))
            if i == 2: D_drug = list(map(lambda x: x + '【D级】' if '-' not in x else '', ele))
            if i == 3: Drug_resistance = ele
        # print(A_drug,C_drug,D_drug,Drug_resistance)
        temp = A_drug + C_drug + D_drug
        temp = list(filter(lambda x: x != "", temp))
        table.cell(4, 2).text = '、'.join(temp) if temp != [] and cell.text!='无' else '无'
        table.cell(5, 2).text = '、'.join(Drug_resistance) if Drug_resistance != [] and Drug_resistance != ['-'] else '无'

        # # 改字体。
        def format_Cell(table, position, size=10, color=RGBColor(0, 0, 0), left=False,no_center_num=None):  # doxc 中 table 和 table 中的cell 都是可变对象。
            for (row, column) in position:
                cell = table.cell(row, column)
                for paragraph in cell.paragraphs:  # 遍历单元格中的段落
                    for run in cell.paragraphs[0].runs:
                        # run = cell.paragraphs[0].runs[0]  # 注意：添加一个新的文本块 add_run() 是不对的，我们这里是修改已有的cell。
                        font = run.font  # 获取字体对象
                        font.name = "Microsoft YaHei"  # 设置字体名称为宋体
                        font.size = Pt(size)  # 设置字体大小为12号
                        font.color.rgb = color  # 设置字体颜色为黑色
                        if left == False:
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 设置段落水平居中
                        elif left == True:
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                        from docx.oxml.ns import qn  # 中文字体。
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # 设置单元格垂直居中

            # 设置除了 第 column_index 列 ，从第二行起非居中对齐。
            def set_column_alignment(table, column_index):
                for i, row in enumerate(table.rows):
                    if i > 0:
                        cell = row.cells[column_index]
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            set_column_alignment(table, no_center_num) if no_center_num != None else print('')
            return table
        format_Cell(table, position=[(2, 1), (3, 1), (4, 2), (5, 2)], size=10, left=True)
        table.cell(2, 1).paragraphs[0].runs[0].text = table.cell(2, 1).paragraphs[0].runs[0].text
        table.cell(2, 1).paragraphs[0].runs[2].font.color.rgb = RGBColor(255, 0, 0)  # 由于前面被集体调整格式了，这里需要重新标红一下。
        table.cell(2, 1).paragraphs[0].paragraph_format.left_indent = Cm(0.2)
        table.cell(3, 1).paragraphs[0].paragraph_format.left_indent = Cm(0.2)
        return self.doc

    @Tools_Decorator(tool='move_table_after', paragraph_end_with="靶向药物注释")
    @Tools_Decorator(tool='format_Table', col_width=[3,3,9.92],First_column_colore=False, header_size=11,Border_all=True)
    @Tools_Decorator(tool='format_Cell',position=position, size=10, color=RGBColor(0, 0, 0),no_center_num=2)
    @Tools_Decorator(tool='Merge_cells_by_first_column', column_index=0)  # 指定合并单元格的列号。
    @Tools_Decorator(tool='format_Table', col_width=[3,3,9.92],First_column_colore=False, header_size=11,Border_all=True)
    def Targeted_drug_annotations(self,position=[]):
        # (读前面的靶向药物提示那张表)
        print(self.targeted_Therapy_Tips_df)
        own = self.targeted_Therapy_Tips_df['推荐本癌种药物'].tolist()
        other = self.targeted_Therapy_Tips_df['推荐其他癌种药物'].tolist()
        cli = self.targeted_Therapy_Tips_df['临床试验药物'].tolist()
        telerence = self.targeted_Therapy_Tips_df['潜在耐药药物'].tolist()
        PotentialBeneficialDrugs_list = [ ele for ele in own+other+cli if ele!='/']
        Potentialdrugresistance_list  = [ ele for ele in telerence if ele!='/']
        print(own,other,cli)
        print(PotentialBeneficialDrugs_list)
        temp = []
        for ele in PotentialBeneficialDrugs_list:
            temp += (ele.split('\n'))
        PotentialBeneficialDrugs_list = temp
        print(PotentialBeneficialDrugs_list)
        temp = []
        for ele in Potentialdrugresistance_list:
            temp += (ele.split('\n'))
        Potentialdrugresistance_list = temp
        # table = self.doc.add_table(rows=1, cols=1)
        useful_df = self.DrugApproval[self.DrugApproval['DrugName'].isin(PotentialBeneficialDrugs_list)]
        useful_df_1 = self.DrugApproval[self.DrugApproval['DrugName'].isin(Potentialdrugresistance_list)]
        # for ele in PotentialBeneficialDrugs_list:
        #     ApprovedContent = self.DrugApproval[self.DrugApproval['DrugName']==ele]['ApprovedContent'].tolist()
        #     ApprovedContent = ''.join(ApprovedContent)
        df_ = pd.DataFrame(columns=['药物敏感性', '药物名称', '用药解析'])
        df_.loc[0] = ['潜在获益药物', '无相关药物', '无']
        df_1 = pd.DataFrame(columns=['药物敏感性', '药物名称', '用药解析'])
        df_1.loc[0] = ['潜在耐药药物', '无相关药物', '无']
        df_merged = useful_df.groupby('DrugName')['ApprovedContent'].apply(' '.join).reset_index() if not useful_df.empty else df_
        df_merged_1 = useful_df_1.groupby('DrugName')['ApprovedContent'].apply(' '.join).reset_index() if not useful_df_1.empty else df_1
        df_merged['药物敏感性'] = '潜在获益药物'
        df_merged_1['药物敏感性'] = '潜在耐药药物'
        df_merged = df_merged[['药物敏感性','DrugName','ApprovedContent']] if not useful_df.empty else df_
        df_merged.columns = ['药物敏感性','药物名称','用药解析']
        df_merged_1 = df_merged_1[['药物敏感性','DrugName','ApprovedContent']] if not useful_df_1.empty else df_1
        df_merged_1.columns = ['药物敏感性','药物名称','用药解析']
        df_merged = pd.concat([df_merged,df_merged_1])
        print(df_merged)
        # 添加一个表格
        table = self.doc.add_table(rows=1, cols=3)
        # 设置表头
        for i, column_name in enumerate(df_merged.columns):
            table.cell(0, i).text = column_name
        # 填充表格内容
        for _, row in df_merged.iterrows():
            new_row = table.add_row().cells
            for i, val in enumerate(row):
                new_row[i].text = str(val)
        # # 合并特定单元格。这里是合并第一列。
        # def merge_cells_by_first_column(table,column_index=0):
        #     first_column = table.columns[column_index]
        #     previous_text = first_column.cells[column_index].text
        #     merge_start_index = 0
        #     for i in range(1, len(first_column.cells)):
        #         current_text = first_column.cells[i].text
        #         if current_text != previous_text:
        #             # 合并单元格，并将第一个单元格设置为保留的值
        #             table.cell(merge_start_index, column_index).merge(table.cell(i - 1, column_index))
        #             table.cell(merge_start_index, column_index).text = previous_text
        #             merge_start_index = i
        #         else:
        #             # 将后续相同值的单元格清空
        #             table.cell(i, column_index).text = ""
        #         previous_text = current_text
        #     # 执行最后一个合并操作
        #     table.cell(merge_start_index, column_index).merge(table.cell(len(first_column.cells) - 1, column_index))
        #     table.cell(merge_start_index, column_index).text = previous_text
        # merge_cells_by_first_column(table) # 合并0列的单元格
        # table.cell(2,0).merge(table.cell(3,0))
        temp_position = [(i, j) for i in range(len(table.rows)) for j in range(len(table.columns))]
        position.extend(temp_position)
        return self.doc

    # col_width = [3.5, 2.5, 2.7, 2, 1.5, 1.8, 2]
    @Tools_Decorator(tool='move_table_after', paragraph_end_with="解读来源于 PharmGKB 数据库，详细结果如下：")
    @Tools_Decorator(tool='format_Table', col_width=[2, 2, 3.28, 2, 2, 2, 2.64],First_column_colore=False, Border_all=True)
    @Tools_Decorator(tool='format_Cell', position=position, size=10, color=RGBColor(0, 0, 0))
    @Tools_Decorator(tool='Merge_cells_by_first_column', column_index=0)  # 指定合并单元格的列号。
    @Tools_Decorator(tool='format_Table', col_width=[2, 2, 3.28, 2, 2, 2, 2.64],First_column_colore=False, Border_all=True)
    def Chemotherapy_drug_testing_1(self,position=[]):
        self.raw_Q80_chemo = self.Q80_chemo
        # 因为 《Chemotherapy_drug_testing_0》表中加了 删除低可信度冲突的内容，所以这里需要提前 (疗效，毒性)删除冲突的低可信度的行。
        # 先加一列 是否是 最高 证据等级。
        from collections import OrderedDict
        def remove_duplicates(seq):
            # 使用OrderedDict去除重复元素并保持顺序
            return list(OrderedDict.fromkeys(seq))
        def custom_sort(data):   # 自定义的证据等级排序。
            def extract_digits_and_letters(element):
                digits = int(re.findall(r'\d+', element)[0])
                letters = ''.join(re.findall(r'[A-Za-z]', element))
                return (digits, letters)
            return sorted(data, key=extract_digits_and_letters, reverse=False)
        # 存储处理后的 drug_df
        processed_dfs = []
        drug_list = remove_duplicates(self.Q80_chemo['药物名称'].tolist())
        grouped = self.Q80_chemo.groupby('药物名称')
        for i, drug in enumerate(drug_list):
            drug_df = grouped.get_group(drug)
            highest_level=custom_sort(drug_df['证据等级'].tolist())[0]
            # print('这种药物的最高证据等级为：',drug,highest_level)
            drug_df['是否是最高证据等级'] = False
            drug_df.loc[drug_df['证据等级'] == highest_level, '是否是最高证据等级'] = True
            drug_df.loc[drug_df['证据等级'] != highest_level, '是否是最高证据等级'] = False
            # print(drug_df[drug_df['证据等级']==highest_level])
            # 解决冲突(疗效和毒性分开处理)
            drug_high_df = drug_df[(drug_df['类别'].str.contains('疗效'))&(drug_df['是否是最高证据等级']==True)]
            # print(drug_high_df)
            matching_strings = drug_high_df['用药提示'].str.extract('(疗效[\u4e00-\u9fff]{2})', expand=False)
            # print(matching_strings)
            is_unique = len(matching_strings.value_counts()) == 1
            if not drug_high_df.empty and is_unique==True:
                # 去掉疗效 冲突的行。
                # print('最高等级 疗效 是唯一的。')
                # 如果要删除的 行不存在，就不用处理。
                drug_df = drug_df[~((drug_df['是否是最高证据等级'] == False) & (drug_df['类别'].str.contains('疗效')) & ~(drug_df['用药提示'].str.contains(drug_high_df['用药提示'].iloc[0])))]
            elif not drug_high_df.empty and is_unique==False:
                # 先统一疗效\毒性 ，再 去掉 低等级冲突的行。
                # print('最高等级 疗效不是唯一的，先统一疗效，再去掉低等级冲突的行。')
                drug_high_df['用药提示'] = drug_high_df['用药提示'].str.replace('疗效[\u4e00-\u9fff]{2}', '疗效减弱', regex=True)
                drug_df.loc[(drug_df['是否是最高证据等级'] == True) & (drug_df['类别'].str.contains('疗效')), '用药提示'] \
                    = drug_df.loc[(drug_df['是否是最高证据等级'] == True) & (drug_df['类别'].str.contains('疗效')), '用药提示'].str.replace('疗效([\u4e00-\u9fff]{2})', '疗效减弱', regex=True)
            drug_high_df = drug_df[(drug_df['类别'].str.contains('毒性')) & (drug_df['是否是最高证据等级'] == True)]
            matching_strings = drug_high_df['用药提示'].str.extract('(毒性[\u4e00-\u9fff]{2})', expand=False)
            is_unique = len(matching_strings.value_counts()) == 1
            if not drug_high_df.empty and is_unique==True:
                # 去掉毒性 冲突的行。
                # print('最高等级 毒性 是唯一的。')
                drug_df = drug_df[~((drug_df['是否是最高证据等级'] == False) & (drug_df['类别'].str.contains('毒性')) & ~(drug_df['用药提示'].str.contains(drug_high_df['用药提示'].iloc[0])))]
            elif not drug_high_df.empty and is_unique==False:
                # 先统一疗效\毒性 ，再 去掉 低等级冲突的行。
                drug_high_df['用药提示'] = drug_high_df['用药提示'].str.replace('毒性[\u4e00-\u9fff]{2}', '毒性增强', regex=True)
                drug_df.loc[(drug_df['是否是最高证据等级'] == True) & (drug_df['类别'].str.contains('毒性')), '用药提示'] \
                    = drug_df.loc[(drug_df['是否是最高证据等级'] == True) & (drug_df['类别'].str.contains('毒性')), '用药提示'].str.replace('毒性([\u4e00-\u9fff]{2})', '毒性增强', regex=True)
            processed_dfs.append(drug_df)

        self.Q80_chemo_complete = pd.concat(processed_dfs)
        # self.Q80_chemo = self.Q80_chemo_complete.drop('是否是最高证据等级', axis=1) # 这一句打开就是 中间表显示出来 写进word。
        self.Q80_chemo = self.raw_Q80_chemo
        # 添加一个表格
        table = self.doc.add_table(rows=1, cols=len(self.Q80_chemo.columns))
        # 设置表头
        for i, column_name in enumerate(self.Q80_chemo.columns):
            table.cell(0, i).text = column_name
        # 填充表格内容
        for _, row in self.Q80_chemo.iterrows():
            new_row = table.add_row().cells
            for i, val in enumerate(row):
                new_row[i].text = str(val)
        temp_position = [(i, j) for i in range(len(table.rows)) for j in range(len(table.columns))]
        position.extend(temp_position)
        return self.doc

    @Tools_Decorator(tool='move_table_after', paragraph_end_with="chemotable")
    @Tools_Decorator(tool='format_Table', col_width=[2,3,2.96],First_column_colore=False, Border_all=True)
    @Tools_Decorator(tool='Specify_cell_color_change', search_text='推荐选用', color=RGBColor(0, 176, 80))
    @Tools_Decorator(tool='Specify_cell_color_change', search_text='谨慎选用',color=RGBColor(255, 0, 0) )
    @Tools_Decorator(tool='format_Cell', position=position, size=10, color=RGBColor(0, 0, 0))
    @Tools_Decorator(tool='Merge_cells_by_first_column', column_index=0)  # 指定合并单元格的列号。
    @Tools_Decorator(tool='format_Table', col_width=[2,3,2.96],First_column_colore=False, Border_all=True)
    def Chemotherapy_drug_testing_0(self,position=[]):
        from collections import OrderedDict
        def remove_duplicates(seq):
            # 使用OrderedDict去除重复元素并保持顺序
            return list(OrderedDict.fromkeys(seq))
        drug_list = remove_duplicates(self.Q80_chemo['药物名称'].tolist())
        # print(drug_list)
        # 这个由 《◆化疗药物检测结果详情及解读》 到 《化疗药物指标及提示》 的逻辑一定要讲清楚：
        # 先在 化药表 中 加一列标注出 是否是最高证据等级。
        # 毒性和 疗效先分两条路 判断，再 联系起来综合判断；
        # 分开判断的逻辑是(这一步发生在Chemotherapy_drug_testing_1中)：
        # 如果低证据等级和高证据等级有冲突，以高证据等级的判读。低证据等级的那行删掉。
        # 同级高证据等级有冲突的，是疗效路则药效减弱，毒性路则毒性增强,表内改动。
        # 综合判断逻辑(这一步发生在Chemotherapy_drug_testing_0中)：
        # （疗效增强，毒性减弱）推荐；（疗效减弱，毒性增强）谨慎；
        # 其余像（疗效减弱，毒性减弱）（疗效减弱，剂量增加）等等把，只要是两个没有同时都满足的，全部被视为 常规推荐；

        # 药物类型对应关系：
        def relationship_of_drug_types(drug: str)-> str:
            options = {
                '环磷酰胺':'烷化剂','替莫唑胺':'烷化剂',
                '氟尿嘧啶':'抗代谢药','替加氟':'抗代谢药','吉西他滨':'抗代谢药','培美曲塞':'抗代谢药','甲氨蝶呤':'抗代谢药',
                '卡培他滨': '抗代谢药','甲氨蝶呤': '抗代谢药',
                '伊立替康':'植物来源抗肿瘤药','长春瑞滨':'植物来源抗肿瘤药','紫杉醇':'植物来源抗肿瘤药','多西他赛':'植物来源抗肿瘤药',
                '长春新碱': '植物来源抗肿瘤药','依托泊苷': '植物来源抗肿瘤药',
                '蒽环类':'抗生素','表柔比星':'抗生素','多柔比星':'抗生素','丝裂霉素':'抗生素','红霉素':'抗生素',
                '他莫昔芬':'激素药','来曲唑':'激素药','阿那曲唑':'激素药',
                '顺铂':'其他','卡铂':'其他','奥沙利铂':'其他','甲酰四氢叶酸':'其他',
                '昂丹司琼':'辅助治疗','格拉司琼':'辅助治疗'
            }
            return options.get(drug, f'unknow {drug}' )
        # 以下逻辑。
        grouped = self.Q80_chemo_complete[self.Q80_chemo_complete['是否是最高证据等级']==True].groupby('药物名称') # 只按照最高证级判断。
        df = pd.DataFrame(columns=['药物类型','药物名称','用药提示'])
        for i,drug in enumerate(drug_list):
            temp_Tip = '常规选用'
            drug_df = grouped.get_group(drug)
            if not drug_df[drug_df['用药提示'].str.contains('疗效增强')].empty and not drug_df[drug_df['用药提示'].str.contains('毒性减弱')].empty:
                temp_Tip = '推荐选用'
            if not drug_df[drug_df['用药提示'].str.contains('疗效减弱')].empty and not drug_df[drug_df['用药提示'].str.contains('毒性增强')].empty:
                temp_Tip = '谨慎选用'
            df.loc[i] = [relationship_of_drug_types(drug), drug, temp_Tip]
        # 以上逻辑。

        # 这个存好，后面填写表用。
        self.danger_chemo_drug_list = df[df['用药提示']=='谨慎选用']['药物名称'].tolist()
        self.recommend_chemo_drug_list = df[df['用药提示']=='推荐选用']['药物名称'].tolist()
        self.common_chemo_drug_list = df[df['用药提示'] == '常规选用']['药物名称'].tolist()
        # 添加一个表格
        table = self.doc.add_table(rows=1, cols=len(df.columns))
        # 设置表头
        for i, column_name in enumerate(df.columns):
            table.cell(0, i).text = column_name
        # 填充表格内容
        for _, row in df.iterrows():
            new_row = table.add_row().cells
            for i, val in enumerate(row):
                new_row[i].text = str(val)

        # 获取表格的总行数
        total_rows = len(table.rows)
        # 检查总行数是否为奇数
        if total_rows % 2 == 0:
            # 添加新行
            new_row = table.add_row().cells
            # 针对新添加的行，为每个单元格设置内容为 "-"
            for cell in new_row:
                cell.text = '-'
        # # 以下为分栏 测试
        target_section = section = self.doc.sections[2]  # 在母板中chemotable前后加上两个 分节符。
        sectPr = target_section._sectPr
        cols = sectPr.xpath('./w:cols')[0]
        cols.set(qn('w:num'), '2')
        cols.set(qn('w:space'), '0')
        # # 以上为分栏 测试

        temp_position = [(i, j) for i in range(len(table.rows)) for j in range(len(table.columns))]
        position.extend(temp_position)
        return self.doc

    @Tools_Decorator(tool='format_Cell', position=position, table_index=3,size=10, color=RGBColor(0, 0, 0),left=True)
    def Evaluation_of_therapeutic_effect_of_chemotherapy_drugs(self,position=[]):
        table = self.doc.tables[3]
        table.cell(2,1).text = '、'.join(self.recommend_chemo_drug_list) if self.recommend_chemo_drug_list!=[] else '无'
        table.cell(3, 1).text = '、'.join(self.common_chemo_drug_list) if self.common_chemo_drug_list!=[] else '无'
        table.cell(4, 1).text = '、'.join(self.danger_chemo_drug_list) if self.danger_chemo_drug_list!=[] else '无'
        position.extend([(2,1),(3, 1),(4, 1)])
        table.cell(2, 1).paragraphs[0].paragraph_format.left_indent = Cm(0.2)
        table.cell(3, 1).paragraphs[0].paragraph_format.left_indent = Cm(0.2)
        table.cell(4, 1).paragraphs[0].paragraph_format.left_indent = Cm(0.2)
        return self.doc

    @Tools_Decorator(tool='move_table_after', paragraph_end_with="通过高通量测序法检测微卫星不稳定（MSI），评估免疫检查点抑制剂获益情况。")
    @Tools_Decorator(tool='format_Table', col_width=[4.64, 2.5, 2.5, 2.78, 3.5], Border_all=False)
    @Tools_Decorator(tool='format_Cell', position=position, size=10, color=RGBColor(0, 0, 0))
    def MSI_detection(self,position=[]):
        table = self.doc.add_table(rows=2, cols=5)
        for i,head in enumerate(['检测项目','检测结果','参考值(%)','检测值(%)','获益情况']):
            table.cell(0,i).text = head
        table.cell(1,0).text = '微卫星不稳定(MSI)'
        # print(self.Q80_msi['percent'].iloc[0],type(self.Q80_msi['percent'].iloc[0]))
        table.cell(1,1).text = 'MSI-L' if float(self.Q80_msi['percent'].iloc[0])<30 else 'MSI-H'
        table.cell(1, 2).text = '30'
        table.cell(1, 3).text = str(self.Q80_msi['percent'].iloc[0])
        table.cell(1, 4).text = '获益一般' if table.cell(1,1).text == 'MSI-L' else '获益较高'
        temp_position = [(i, j) for i in range(len(table.rows)) for j in range(len(table.columns))]
        position.extend(temp_position)
        return self.doc

    @Tools_Decorator(tool='format_Cell', table_index=2,position=position, size=10, color=RGBColor(0, 0, 0))
    def Evaluation_of_the_therapeutic_effect_of_immunotherapy(self,position=[]):
        for temp_table in self.doc.tables:
            if len(temp_table.columns) == 3 and temp_table.cell(0, 0).text == '免疫药治疗疗效评估':
                target_table = temp_table
                break
        for temp_table in self.doc.tables:
            if len(temp_table.columns) == 5 and temp_table.cell(0, 0).text == '检测项目'\
                    and temp_table.cell(0, 1).text == '检测结果'\
                    and temp_table.cell(0, 2).text == '参考值(%)'\
                    and temp_table.cell(0, 3).text == '检测值(%)'\
                    and temp_table.cell(0, 4).text == '获益情况':
                target_table_ref = temp_table
                break
        table = target_table
        table_ref = target_table_ref
        if self.clinname=='肠癌':
            if 'MSI-L' in table_ref.cell(1, 1).text:
                table.cell(2,1).text = ' MSI-L（微卫星不稳定度较低）'
                table.cell(2,2).text = '结直肠癌患者可能预后较差\n结直肠癌患者对5-FU方案可能敏感\n免疫检查点抑制剂获益一般\n患林奇综合征风险较低'
            else:
                table.cell(2, 1).text = ' MSI-H（微卫星不稳定度较高）'
                table.cell(2, 2).text = '结直肠癌患者可能预后较好\n结直肠癌患者对5-FU方案可能不敏感\n免疫检查点抑制剂获益较好\n患林奇综合征风险较高'
        else:
            if 'MSI-L' in table_ref.cell(1, 1).text:
                table.cell(2,1).text = ' MSI-L（微卫星不稳定度较低）'
                table.cell(2,2).text = '免疫检查点抑制剂获益一般\n患林奇综合征风险较低'
            else:
                table.cell(2, 1).text = ' MSI-H（微卫星不稳定度较高）'
                table.cell(2, 2).text = '免疫检查点抑制剂获益较好\n患林奇综合征风险较高'
        position.extend([(2,1),(2,2)])
        return self.doc

    @Tools_Decorator(tool='move_table_after', paragraph_end_with="多靶点药物和联合用药方案评估")
    @Tools_Decorator(tool='Specify_cell_color_change', search_text='慎选', color=RGBColor(255, 0, 0))
    @Tools_Decorator(tool='format_Table', col_width=[1.2, 4.32, 4, 1.2, 4, 1.2], First_column_colore=False, Border_all=True,no_colore_column=[0,1])
    @Tools_Decorator(tool='format_Cell', position=position, size=10, color=RGBColor(0, 0, 0))
    @Tools_Decorator(tool='Merge_cells_by_first_column', column_index=1)  # 指定合并单元格的列号。
    @Tools_Decorator(tool='Merge_cells_by_first_column', column_index=0)  # 指定合并单元格的列号。
    @Tools_Decorator(tool='format_Table', col_width=[1.2, 4.32, 4, 1.2, 4, 1.2], First_column_colore=False,Border_all=True,no_colore_column=[0,1])
    def DrugCombinedPlan_func(self,position=[]):
        print(self.cancer,'是癌症种类。')
        df = self.DrugCombinedPlan[self.DrugCombinedPlan['癌种'].apply(lambda x: x in self.cancer)]
        # print(df)
        print('突变过滤后数量是 ',len(self.cell),[(ele.gene,ele.amino_acid) for ele in self.cell])
        def judge_wheter_optional(List_: List[str])-> str:
            if List_ == ['']:
                return '可选'
            else:
                mutation_gene_str = ','.join([ele.gene for ele in self.cell])
                for obj in List_:
                    if obj in self.danger_chemo_drug_list:
                        return '慎选'
                    if hasattr(self, 'MSI_result'):
                        print("变量 self.MSI_result 存在，现在看检测结果。")
                        if 'MSI-H' in self.MSI_result and obj == 'MSI-H':
                            return '慎选'
                    if hasattr(self, 'MMR_result'):
                        print("变量 self.MMR_result 存在，现在看检测结果。")
                        if self.MMR_result!='未检出' and obj == 'dMMR':
                            return '慎选'
                    if obj in mutation_gene_str:
                        return '慎选'
            return '可选'
        df.fillna('', inplace=True)
        df['评估结果'] = ''
        for index,row in df.iterrows():
            print(type(row['评估内容']),row['评估内容'])
            Evaluation_Content_List = re.split(r"[,.;，/]",row['评估内容'])
            print(Evaluation_Content_List)
            df.at[index, '评估结果'] = judge_wheter_optional(Evaluation_Content_List)
        print(df)
        # 添加一个表格
        table = self.doc.add_table(rows=1, cols=6)
        # 设置表头
        for i, column_name in enumerate(['方案用途','使用范围','用药方案','方案等级','药物机制','临床提示']):
            table.cell(0, i).text = column_name
        # 填充表格内容
        for _, row in df.iterrows():
            new_row = table.add_row().cells
            new_row[0].text = str(row['方案用途'])
            new_row[1].text = str(row['使用范围'])
            new_row[2].text = str(row['用药方案'])
            new_row[3].text = str(row['方案等级']).replace('类','')
            new_row[4].text = str(row['药物机制'])
            new_row[5].text = str(row['评估结果'])
        temp_position = [(i, j) for i in range(len(table.rows)) for j in range(len(table.columns))]
        position.extend(temp_position)
        return self.doc

    def add_header(self, manual_Date=None):  # 改页眉。
        section = self.doc.sections[1]
        header = section.header
        header.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中显示
        header.tables[0].cell(0, 0).paragraphs[0].text = '上海解码医学检验所'  # '清港泉生物科技有限公司 '
        data_ = f'{datetime.datetime.now().date().strftime("%Y-%m-%d")}' if manual_Date == None else manual_Date
        sample_id, sample_name = self.BC17_meta['id'].iloc[0], self.BC17_meta['name'].iloc[0]
        header.tables[0].cell(0, 1).paragraphs[0].text = sample_id.replace('-T', '').replace('-N',
                                                                                             '') + ' | ' + sample_name + ' | ' + data_
        header.tables[0].cell(0, 0).paragraphs[0].runs[0].font.color.rgb = RGBColor(53, 105, 155)
        header.tables[0].cell(0, 1).paragraphs[0].runs[0].font.color.rgb = RGBColor(53, 105, 155)
        return self.doc

    def add_date(self, manual_Date=None):
        # 通过匹配table的表头来定位table。
        for table in self.doc.tables:
            if len(table.columns) == 2 and table.cell(0, 0).paragraphs[0].text == '检测人：' and \
                    table.cell(0, 1).paragraphs[0].text == '复核人：':
                target_table = table
        target_table.cell(1,
                          0).text = f'报告日期：{datetime.datetime.now().date().strftime("%Y-%m-%d")}' if manual_Date == None else f'报告日期：{manual_Date}'
        from docx.oxml.ns import qn  # 中文字体。
        cell = target_table.cell(1, 0)
        for paragraph in cell.paragraphs:  # 遍历单元格中的段落
            run = cell.paragraphs[0].runs[0]  # 注意：添加一个新的文本块 add_run() 是不对的，我们这里是修改已有的cell。
            font = run.font  # 获取字体对象
            font.name = "Microsoft YaHei"  # 设置字体名称
            from docx.oxml.ns import qn  # 中文字体。
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        return self.doc

    def delete_rubbish(self):
        for ele in ['chemotable','Somatic_variation_results_table','drugtable','pathogenictable','unknowntable','msitable','chemodetailtable']:
            try:
                target = self.class_Tools.fixed_position(self.doc, End_character=ele)
                if target!=None:
                    self.class_Tools.delete_paragraph(target)
            except:
                print('无此段落。')
                continue
        return self.doc

if __name__ == '__main__':
    # 2022WSSW005892-T.summary.xlsx 2023WSSW001035-T.summary.xlsx 2023WSSW001754-T
    # sample = '2023WSSW001328-T'
    # date = '230626'
    sample = sys.argv[1]
    date = sys.argv[2]
    word_name = sys.argv[3]
    print('sample',sample,'date',date,'word_name',word_name)
    q80 = Q80(f'/archive/{date}/{sample}.summary.xlsx') # 2023WSSW001612-T.summary.xlsx 2023WSSW001320-T.summary.xlsx 2023WSSW000786-T.summary.xlsx
    doc = q80.sample_information(position=[])
    doc = q80.Details_of_genetic_testing_results(position=[],position_1=[]) # 这一步其实对self.cell做了一次过滤。
    doc = q80.Somatic_variation_results_0(position=[])
    doc = q80.Somatic_variation_results_1(position=[])
    doc = q80.targeted_Therapy_Tips(position=[])
    doc = q80.Analysis_of_Somatic_Variant_Genes_and_Loci(position=[])
    doc = q80.quality_Control_Results(position=[])
    doc = q80.machining_table_1()
    doc = q80.Chemotherapy_drug_testing_1(position=[])
    doc = q80.Chemotherapy_drug_testing_0(position=[])
    doc = q80.Evaluation_of_therapeutic_effect_of_chemotherapy_drugs( position=[])
    doc = q80.MSI_detection(position=[])
    doc = q80.Evaluation_of_the_therapeutic_effect_of_immunotherapy(position=[])
    doc = q80.Targeted_drug_annotations(position=[])
    doc = q80.DrugCombinedPlan_func(position=[])

    doc = q80.add_header()
    doc = q80.add_date()
    doc = q80.delete_rubbish()

    folder_path = f'/archive/word/{date}'
    if not os.path.exists(folder_path):
        os.makedirs(folder_path)
    q80.doc.save(f'/archive/word/{date}/{word_name}.docx')


